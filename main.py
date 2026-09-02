"""
Модуль LAIM kriteria selector - выбор критериев разметки.

Основная функция: сопоставление бизнес-метрик из отчёта о разработке
с фактическими столбцами в датасете с использованием LLM.

Функции:
    read_docx: Чтение docx файла с инструкцией по разметке
    make_giga_request: Запрос к GigaChat
    make_sds_request: Запрос к SDS моделям через AI Gateway
    main: Основная функция модуля
"""

import json
import hashlib
import io
import logging
import os
import pickle
import re
import time
from pathlib import Path

import pandas as pd
import requests
from docx import Document

# Повторы обращения к модели. Первая повторная попытка снимает почти все
# единичные сбои шлюза; экспоненциальная задержка даёт квоте восстановиться —
# немедленные повторы приходятся на тот же всплеск нагрузки и сгорают впустую
# (зафиксированный риск ревью контура: ретраи без задержки).
LLM_ATTEMPTS = 3
RETRY_BASE_DELAY_SECONDS = 2.0

# Обрезка длинных значений в примере данных для промта. Реплики диалогов
# занимают тысячи символов; для сопоставления метрики с колонками модели
# достаточно начала значения — оно показывает, текст это или метка.
SAMPLE_VALUE_MAX_CHARS = 200

# Предел длины инструкции разметки в промте. Инструкции пилота — единицы
# страниц (< 15 тыс. символов); документ длиннее — это, как правило, полный
# отчёт, поданный не в тот порт, и тянуть его в промт целиком значит вытеснить
# из контекста сами данные. Обрезка логируется.
INSTRUCTION_MAX_CHARS = 30_000

# Версия контракта между selector и всеми потребителями измерительной ветки.
# Старые поля main_metric/other_metrics сохранены, поэтому assessor и тесты,
# которые о расширении не знают, продолжают работать без изменений.
METRIC_SPEC_VERSION = "metric-spec.v2"
DERIVED_SCORE_COLUMN = "laim_key_metric"
ROW_AGGREGATIONS = ("identity", "majority_vote")
METRIC_STRATEGIES = (
    "mean_score",
    "weighted_accuracy",
    "accuracy",
    "mean_multi_marker",
    "categorical_label",
)

# DataArtifact в SberDS не имеет единственного Python-представления. Файл
# может доехать как локальный путь (часто без расширения), raw bytes или как
# однострочный parquet-контейнер с bytes/path внутри. Сначала нормализуем этот
# транспорт и только затем разбираем внутренний DOCX.
_PARQUET_MAGIC = b"PAR1"
_ZIP_MAGIC = b"PK\x03\x04"
_ARTIFACT_BINARY_KEYS = (
    "bin", "bytes", "content", "data", "payload", "value",
    "unstructured_data",
)
_ARTIFACT_PATH_KEYS = ("path", "local_path", "file", "__file__")
_ARTIFACT_EXTENSION_KEYS = (
    "ext", "extension", "suffix", "filename", "file_name", "name",
)
_ARTIFACT_MAX_UNWRAP_DEPTH = 6


class ArtifactTransportError(ValueError):
    """Поданный DataArtifact невозможно восстановить до исходного файла."""


def _json_bytes(value) -> bytes:
    """Стабильное представление артефакта только для вычисления digest."""
    return json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        default=lambda item: item.item() if hasattr(item, "item") else str(item),
    ).encode("utf-8")


def _dataframe_sha256(frame: pd.DataFrame) -> str:
    """Digest схемы и значений DataFrame без вывода содержимого в контракт.

    Хеширование идёт по колонкам, поэтому не создаёт строковую копию всего
    файла. Для ячеек со списками/словарями используется стабильный JSON.
    """
    digest = hashlib.sha256()
    schema = {
        "n_rows": int(len(frame)),
        "columns": [str(column) for column in frame.columns],
        "dtypes": [str(dtype) for dtype in frame.dtypes],
    }
    digest.update(_json_bytes(schema))
    for position in range(len(frame.columns)):
        series = frame.iloc[:, position]
        try:
            hashed = pd.util.hash_pandas_object(
                series, index=True, categorize=True
            )
        except TypeError:
            normalized = series.map(
                lambda value: _json_bytes(value).decode("utf-8")
            )
            hashed = pd.util.hash_pandas_object(
                normalized, index=True, categorize=True
            )
        digest.update(hashed.to_numpy().tobytes())
    return digest.hexdigest()


def _identity_strings(value):
    """Строковые фрагменты structured output для поиска заявленной identity."""
    if isinstance(value, dict):
        for nested in value.values():
            yield from _identity_strings(nested)
    elif isinstance(value, (list, tuple)):
        for nested in value:
            yield from _identity_strings(nested)
    elif isinstance(value, str):
        yield value


def _artifact_extension(value) -> str:
    """Нормализовать ext либо извлечь суффикс из имени файла."""
    if value in (None, ""):
        return ""
    text = str(value).strip().lower()
    if not text:
        return ""
    suffix = Path(text).suffix
    if suffix:
        return suffix
    return text if text.startswith(".") else f".{text}"


def _artifact_scalar_present(value) -> bool:
    if value is None:
        return False
    if pd.api.types.is_scalar(value):
        try:
            if bool(pd.isna(value)):
                return False
        except (TypeError, ValueError):
            pass
    return not (isinstance(value, str) and not value.strip())


def _artifact_frame_payload(
    frame: pd.DataFrame,
    extension: str,
    depth: int,
) -> tuple[bytes | None, str]:
    """Извлечь единственный payload из parquet-обёртки DataArtifact."""
    if frame.empty:
        raise ArtifactTransportError("parquet-контейнер DataArtifact пуст")

    extension_hint = extension
    for key in _ARTIFACT_EXTENSION_KEYS:
        if key not in frame.columns:
            continue
        values = [value for value in frame[key].tolist()
                  if _artifact_scalar_present(value)]
        if values:
            extension_hint = _artifact_extension(values[0])
            break

    payload_columns = [
        key for key in (*_ARTIFACT_BINARY_KEYS, *_ARTIFACT_PATH_KEYS)
        if key in frame.columns
    ]
    for column in payload_columns:
        values = [value for value in frame[column].tolist()
                  if _artifact_scalar_present(value)]
        if not values:
            continue
        if len(values) != 1:
            raise ArtifactTransportError(
                f"parquet-контейнер содержит {len(values)} payload-значений "
                f"в колонке {column!r}; ожидалось одно"
            )
        return _development_report_bytes(
            values[0], _extension=extension_hint, _depth=depth + 1
        )

    metadata_columns = set(_ARTIFACT_EXTENSION_KEYS)
    cells = []
    for column in frame.columns:
        if column in metadata_columns:
            continue
        cells.extend(
            value for value in frame[column].tolist()
            if _artifact_scalar_present(value)
        )
    if len(cells) != 1:
        raise ArtifactTransportError(
            "parquet-контейнер DataArtifact должен содержать ровно один "
            f"bytes/path payload; найдено {len(cells)}"
        )
    return _development_report_bytes(
        cells[0], _extension=extension_hint, _depth=depth + 1
    )


_DIRECTORY_SERVICE_FILES = ("_SUCCESS", "_started", "_committed")


def _artifact_directory_files(directory: Path) -> list[Path]:
    """Содержательные файлы каталога порта (служебные Spark-маркеры игнорируются).

    Платформа монтирует порт «as files and folders»: рядом с полезным файлом
    лежат _SUCCESS/_committed_*/*.crc — прежнее требование «ровно один файл»
    роняло чтение валидационного артефакта.
    """
    return sorted(
        path for path in directory.rglob("*")
        if path.is_file()
        and "__MACOSX" not in path.parts
        and not path.name.startswith((".", "._"))
        and not path.name.endswith(".crc")
        and not any(path.name.startswith(marker) for marker in _DIRECTORY_SERVICE_FILES)
    )


def _artifact_directory_file(directory: Path) -> Path:
    files = _artifact_directory_files(directory)
    if len(files) != 1:
        raise ArtifactTransportError(
            f"каталог DataArtifact должен содержать ровно один содержательный "
            f"файл; найдено {len(files)}: {directory}"
        )
    return files[0]


def _development_report_bytes(
    value,
    *,
    _extension: str = "",
    _depth: int = 0,
) -> tuple[bytes | None, str]:
    """Восстановить исходные байты отчёта из транспорта DataArtifact.

    Поддержаны raw bytes, локальный путь/каталог, платформенные dict-обёртки,
    DataFrame/Series и parquet bytes с единственным bytes/path payload.
    """
    if value is None:
        return None, _artifact_extension(_extension)
    if _depth > _ARTIFACT_MAX_UNWRAP_DEPTH:
        raise ArtifactTransportError(
            "превышена допустимая глубина вложенности DataArtifact"
        )

    extension = _artifact_extension(_extension)
    if isinstance(value, dict):
        for key in _ARTIFACT_EXTENSION_KEYS:
            if _artifact_scalar_present(value.get(key)):
                extension = _artifact_extension(value[key])
                break
        for key in (*_ARTIFACT_BINARY_KEYS, *_ARTIFACT_PATH_KEYS):
            if key in value and _artifact_scalar_present(value[key]):
                return _development_report_bytes(
                    value[key], _extension=extension, _depth=_depth + 1
                )
        return None, extension

    if isinstance(value, pd.Series):
        known = set(value.index) & set(
            (*_ARTIFACT_BINARY_KEYS, *_ARTIFACT_PATH_KEYS,
             *_ARTIFACT_EXTENSION_KEYS)
        )
        if known:
            return _development_report_bytes(
                value.to_dict(), _extension=extension, _depth=_depth + 1
            )
        value = value.to_frame().T
    if isinstance(value, pd.DataFrame):
        return _artifact_frame_payload(value, extension, _depth)

    if isinstance(value, memoryview):
        value = value.tobytes()
    if isinstance(value, bytearray):
        value = bytes(value)
    if isinstance(value, bytes):
        if value[:4] == _PARQUET_MAGIC:
            try:
                frame = pd.read_parquet(io.BytesIO(value))
            except Exception as error:
                raise ArtifactTransportError(
                    "parquet-контейнер DataArtifact не прочитан: "
                    f"{type(error).__name__}: {error}"
                ) from error
            # .parquet описывает транспортную обёртку, а не внутренний файл.
            inner_extension = "" if extension == ".parquet" else extension
            return _artifact_frame_payload(frame, inner_extension, _depth)
        return value, extension

    if isinstance(value, (str, Path)):
        raw_text = str(value)
        path = Path(value)
        try:
            if path.is_dir():
                parts = _artifact_directory_files(path)
                if len(parts) > 1 and all(
                    part.suffix.lower() == ".parquet" or part.name.startswith("part-")
                    for part in parts
                ):
                    frames = [pd.read_parquet(part) for part in parts]
                    frame = pd.concat(frames, ignore_index=True)
                    return _artifact_frame_payload(frame, "", _depth)
                path = _artifact_directory_file(path)
            if path.is_file():
                path_extension = path.suffix.lower()
                path_bytes = path.read_bytes()
                return _development_report_bytes(
                    path_bytes,
                    _extension=path_extension or extension,
                    _depth=_depth + 1,
                )
        except OSError:
            # Длинный structured text является содержимым отчёта, а не путём.
            # На некоторых ОС попытка stat такого значения даёт ENAMETOOLONG.
            return None, extension

        looks_like_path = (
            isinstance(value, Path)
            or raw_text.startswith(("hdfs://", "viewfs://", "/user/"))
            or "/" in raw_text
            or "\\" in raw_text
            or Path(raw_text).suffix.lower() in (".docx", ".parquet")
        )
        if looks_like_path:
            if raw_text.startswith(("hdfs://", "viewfs://", "/user/")):
                raise ArtifactTransportError(
                    "получен только HDFS URI, а не локально смонтированный "
                    "файл/bytes DataArtifact: " + raw_text
                )
            raise ArtifactTransportError(
                f"локальный путь DataArtifact не найден: {raw_text}"
            )
    return None, extension


def _development_report_identity_text(value) -> str:
    """Текст raw DOCX только для поиска явно записанных идентификаторов."""
    raw, extension = _development_report_bytes(value)
    if raw is None:
        return ""
    if extension in (".docx", "docx", "") or raw[:4] == _ZIP_MAGIC:
        try:
            document = Document(io.BytesIO(raw))
        except Exception:
            if extension in (".docx", "docx"):
                logging.warning(
                    "kriteria-selector: raw отчёт помечен как DOCX, но не распарсен"
                )
                return ""
        else:
            parts = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return ""


def _declared_artifact_identity(development_report) -> dict:
    """Извлечь только явные идентификаторы, не выводя их по имени файла."""
    solution_ids = set()
    model_version_ids = set()
    distributives = set()
    if isinstance(development_report, dict):
        for mapping in _walk_dicts(development_report):
            for key in ("agent_id", "agent_ci", "solution_id"):
                value = mapping.get(key)
                if value not in (None, "", "-"):
                    solution_ids.add(str(value).strip().upper())
            for key in ("model_version_id", "version_id"):
                value = mapping.get(key)
                if value not in (None, "", "-"):
                    model_version_ids.add(str(value).strip())
            value = mapping.get("distributive")
            if value not in (None, "", "-"):
                distributives.add(str(value).strip().upper())

    for text in _identity_strings(development_report):
        solution_ids.update(
            match.upper() for match in re.findall(r"\bCI\d{8}\b", text, re.IGNORECASE)
        )
        distributives.update(
            match.upper() for match in re.findall(
                r"\bD-\d{2}\.\d{3}\.\d{2}-\d+\b", text, re.IGNORECASE
            )
        )
        # В шапке validation report пара записана как
        # ``ID решения / ID версии: 300196 / 342293``. Общий regex без
        # понимания пары принимал первый номер за model_version_id.
        paired_identity = re.compile(
            r"ID\s+GenAI-решения\s*/\s*ID\s+версии\s+GenAI-решения"
            r"\s*[:|]\s*\d+\s*/\s*(\d{5,})",
            re.IGNORECASE,
        )
        model_version_ids.update(paired_identity.findall(text))
        text_without_pair = paired_identity.sub("", text)
        model_version_ids.update(
            re.findall(
                r"ID\s+версии(?:\s+GenAI-решения|\s+модели\s+в\s+БМ)?"
                r"\s*(?:[|:]\s*)?(?:ID\s*)?(\d{5,})",
                text_without_pair,
                re.IGNORECASE,
            )
        )
    return {
        "solution_ids": sorted(solution_ids),
        "model_version_ids": sorted(model_version_ids),
        "distributives": sorted(distributives),
    }


def _expected_run_identity(run_context) -> dict:
    if not isinstance(run_context, dict):
        return {}
    # Источник selection может завернуть контекст в selection/run_context.
    for key in ("selection", "run_context"):
        if isinstance(run_context.get(key), dict) and "agent_ci" not in run_context:
            run_context = run_context[key]
    expected = {}
    agent = run_context.get("agent_ci") or run_context.get("agent_id")
    if agent not in (None, "", "-"):
        expected["agent_ci"] = str(agent).strip().upper()
    distributive = run_context.get("distributive")
    if distributive not in (None, "", "-"):
        raw = str(distributive).strip().upper()
        match = re.search(r"D-\d{2}\.\d{3}\.\d{2}-\d+", raw)
        expected["distributive"] = match.group(0) if match else raw
    model_version = run_context.get("model_version_id") or run_context.get("version_id")
    if model_version not in (None, "", "-"):
        expected["model_version_id"] = str(model_version).strip()
    return expected


def _compare_artifact_identity(declared: dict, expected: dict) -> tuple[str, list, list]:
    if not expected:
        has_declared = any(declared.values())
        return ("declared_unverified" if has_declared else "not_provided", [], [])
    verified_fields = []
    mismatches = []
    checks = (
        ("agent_ci", "solution_ids"),
        ("distributive", "distributives"),
        ("model_version_id", "model_version_ids"),
    )
    for expected_key, declared_key in checks:
        expected_value = expected.get(expected_key)
        declared_values = declared.get(declared_key) or []
        if not expected_value or not declared_values:
            continue
        if expected_value in declared_values:
            verified_fields.append(expected_key)
        else:
            mismatches.append(
                f"{expected_key}={expected_value!r} не найден среди {declared_values!r}"
            )
    if mismatches:
        return "mismatch", verified_fields, mismatches
    if verified_fields:
        return "verified", verified_fields, []
    return "expected_only", [], []


def _artifact_provenance(
    frame: pd.DataFrame,
    instruction_text: str | None,
    development_report,
    run_context=None,
    development_report_artifact=None,
    validation_report_artifact=None,
) -> dict:
    """Минимальный паспорт реально использованных входных артефактов.

    В отчёт уходят только digest, схема и размеры. Тексты запросов, ответов и
    инструкции не дублируются: это одновременно ограничивает утечку данных и
    позволяет доказать, на каких именно байтах был построен MetricSpec.
    """
    instruction = str(instruction_text or "")
    report_supplied = development_report is not None
    report_digest = (
        hashlib.sha256(_json_bytes(development_report)).hexdigest()
        if report_supplied else None
    )
    artifact_bytes, _ = _development_report_bytes(development_report_artifact)
    validation_bytes, _ = _development_report_bytes(validation_report_artifact)
    identity_inputs = [
        development_report,
        _development_report_identity_text(development_report_artifact),
    ]
    declared_parts = [
        _declared_artifact_identity(value) for value in identity_inputs
    ]
    declared_identity = {
        key: sorted({item for part in declared_parts for item in part[key]})
        for key in ("solution_ids", "model_version_ids", "distributives")
    }
    expected_identity = _expected_run_identity(run_context)
    identity_status, verified_fields, identity_mismatches = _compare_artifact_identity(
        declared_identity, expected_identity
    )
    validation_text = _development_report_identity_text(validation_report_artifact)
    validation_declared = _declared_artifact_identity(validation_text)
    validation_expected = {
        key: expected_identity[key]
        for key in ("agent_ci", "model_version_id")
        if expected_identity.get(key)
    }
    validation_identity_status, validation_verified, validation_mismatches = (
        _compare_artifact_identity(
            {**validation_declared, "distributives": []},
            validation_expected,
        )
    )
    return {
        "basket": {
            "sha256": _dataframe_sha256(frame),
            "n_rows": int(len(frame)),
            "columns": [str(column) for column in frame.columns],
            "dtypes": {str(column): str(dtype)
                       for column, dtype in zip(frame.columns, frame.dtypes)},
        },
        "instruction": {
            "provided": bool(instruction),
            "sha256": (
                hashlib.sha256(instruction.encode("utf-8")).hexdigest()
                if instruction else None
            ),
            "n_chars": len(instruction),
        },
        "development_report": {
            "provided": report_supplied or artifact_bytes is not None,
            "sha256": report_digest,
            "artifact_sha256": (
                hashlib.sha256(artifact_bytes).hexdigest()
                if artifact_bytes is not None else None
            ),
            "artifact_n_bytes": len(artifact_bytes) if artifact_bytes is not None else 0,
        },
        "validation_report": {
            "provided": validation_bytes is not None,
            "artifact_sha256": (
                hashlib.sha256(validation_bytes).hexdigest()
                if validation_bytes is not None else None
            ),
            "artifact_n_bytes": (
                len(validation_bytes) if validation_bytes is not None else 0
            ),
            "declared_solution_ids": validation_declared["solution_ids"],
            "identity_status": validation_identity_status,
            "identity_verified_fields": validation_verified,
            "identity_mismatches": validation_mismatches,
        },
        "declared_identity": declared_identity,
        "expected_identity": expected_identity,
        "identity_status": identity_status,
        "identity_verified_fields": verified_fields,
        "identity_mismatches": identity_mismatches,
    }


_VALIDATION_LABEL_HINTS = (
    "accuracy", "metric", "score", "оцен", "метрик", "точност",
    "релевант", "полнот", "классификац", "итог",
)
_LABEL_HEADERS = (
    "наименование критерия", "критерий", "метрика", "metric", "metric name",
)
_VALUE_HEADERS = ("значение", "value", "metric value")


def _normalize_text(value) -> str:
    return " ".join(str(value or "").strip().lower().replace("ё", "е").split())


def _parse_report_number(value) -> tuple[float | None, float | None]:
    """Число и допуск, заданный точностью публикации в отчёте.

    Допуск равен единице последнего знака: часть отчётов пилота не округляет,
    а усекает число (0.973684 записано как 0.9736). Это всё ещё существенно
    строже допуска в процентный пункт и не позволяет подобрать произвольную
    формулу под опубликованный результат.
    """
    text = str(value or "").strip().replace(" ", "").replace(",", ".")
    match = re.fullmatch(r"([+-]?\d+(?:\.\d+)?)\s*(%)?", text)
    if not match:
        return None, None
    number = float(match.group(1))
    decimals = len(match.group(1).partition(".")[2])
    tolerance = 10.0 ** (-decimals) if decimals else 1.0
    if match.group(2):
        number /= 100.0
        tolerance /= 100.0
    return number, tolerance + 1e-12


def _validation_report_content(value) -> tuple[str, list[list[list[str]]], str | None]:
    """Распаковать DataArtifact и прочитать внутренний отчёт без LLM."""
    try:
        raw, extension = _development_report_bytes(value)
    except ArtifactTransportError as error:
        return "", [], str(error)
    if raw is None:
        return "", [], "артефакт не содержит доступных байтов"
    # DOCX — zip-контейнер. Сигнатура важнее расширения: локальный путь порта
    # в контуре часто не сохраняет исходное имя файла.
    if extension in (".docx", "docx", "") or raw[:4] == _ZIP_MAGIC:
        try:
            document = Document(io.BytesIO(raw))
        except Exception as error:
            if extension in (".docx", "docx") or raw[:4] == _ZIP_MAGIC:
                return "", [], f"DOCX не прочитан: {type(error).__name__}"
        else:
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            tables = []
            for table in document.tables:
                rows = [[cell.text for cell in row.cells] for row in table.rows]
                tables.append(rows)
                paragraphs.extend(" | ".join(row) for row in rows)
            return "\n".join(paragraphs), tables, None
    try:
        return raw.decode("utf-8"), [], None
    except UnicodeDecodeError:
        return "", [], "формат отчёта не поддержан"


def _validation_observations(tables) -> list[dict]:
    """Извлечь только значения из типизированных metric/value-таблиц.

    Нижние доверительные границы, корреляции и p-value не принимаются за
    baseline: колонка обязана иметь точный заголовок ``Значение``/``Value``.
    """
    observations = []
    for table_index, rows in enumerate(tables):
        if not rows:
            continue
        headers = [_normalize_text(cell) for cell in rows[0]]
        label_index = next(
            (index for index, header in enumerate(headers)
             if header in _LABEL_HEADERS),
            None,
        )
        value_index = next(
            (index for index, header in enumerate(headers)
             if header in _VALUE_HEADERS),
            None,
        )
        if label_index is None or value_index is None:
            continue
        for row_index, row in enumerate(rows[1:], start=1):
            if max(label_index, value_index) >= len(row):
                continue
            label = " ".join(str(row[label_index]).split())
            normalized_label = _normalize_text(label)
            if not label or not any(
                hint in normalized_label for hint in _VALIDATION_LABEL_HINTS
            ):
                continue
            value, tolerance = _parse_report_number(row[value_index])
            if value is None:
                continue
            aggregation = (
                "weighted" if any(key in normalized_label for key in (
                    "weighted", "взвеш",
                )) else
                "macro" if "macro" in normalized_label else
                "unspecified"
            )
            observations.append({
                "label": label,
                "value": value,
                "tolerance": tolerance,
                "aggregation": aggregation,
                "table_index": table_index,
                "row_index": row_index,
            })
    return observations


def _validation_evidence(value, df: pd.DataFrame) -> dict:
    if value is None:
        return {
            "provided": False,
            "status": "not_provided",
            "observations": [],
            "majority_vote_cue": False,
            "mentioned_columns": [],
        }
    text, tables, error = _validation_report_content(value)
    if error:
        raise ArtifactTransportError(
            "validation_report_artifact не прочитан: " + error
        )
    normalized = _normalize_text(text)
    majority = bool(
        re.search(r"большинств\w* голос", normalized)
        or re.search(
            r"(?:итогов\w* результат|финальн\w* метк).{0,160}большинств",
            normalized,
        )
        or "majority vote" in normalized
        or "majority voting" in normalized
    )
    mentioned = []
    for column in df.columns:
        name = str(column).strip()
        if not name:
            continue
        # Границы слова не дают колонке mark совпасть с mark1/mark2.
        if re.search(rf"(?<![\w]){re.escape(name)}(?![\w])", text, re.IGNORECASE):
            mentioned.append(name)
    observations = _validation_observations(tables)
    return {
        "provided": True,
        "status": "parsed",
        "observations": observations,
        "majority_vote_cue": majority,
        "mentioned_columns": mentioned,
        "text_for_prompt": text[:12_000],
    }


def _strict_majority(frame: pd.DataFrame, source_columns: list[str]) -> pd.Series:
    votes = frame[source_columns].apply(pd.to_numeric, errors="coerce")
    modes = votes.mode(axis=1, dropna=True)
    result = pd.Series(float("nan"), index=frame.index, dtype=float)
    if modes.empty:
        return result
    candidate = modes.iloc[:, 0]
    count = votes.eq(candidate, axis=0).sum(axis=1)
    available = votes.notna().sum(axis=1)
    valid = candidate.notna() & (count > available / 2.0)
    result.loc[valid] = candidate.loc[valid].astype(float)
    return result


def _collapse_reference_groups(frame: pd.DataFrame, columns: list[str]) -> pd.DataFrame:
    """Вернуть по одному наблюдению на единицу оценки из исходной корзины."""
    if "reference_group_id" not in frame.columns:
        return frame
    usable = [column for column in columns if column and column in frame.columns]
    if not frame["reference_group_id"].duplicated().any() or not usable:
        return frame
    grouped = frame.groupby("reference_group_id", dropna=False)[usable].nunique(
        dropna=True
    )
    if (grouped > 1).any().any():
        return frame
    return frame.drop_duplicates("reference_group_id")


def _measured_value(
    frame: pd.DataFrame, score_column: str, weight_column: str | None = None,
) -> float | None:
    columns = [score_column] + ([weight_column] if weight_column else [])
    measured = _collapse_reference_groups(frame, columns)
    score = pd.to_numeric(measured[score_column], errors="coerce")
    if weight_column:
        weight = pd.to_numeric(measured[weight_column], errors="coerce")
        valid = score.notna() & weight.notna() & (weight > 0)
        if not valid.any() or float(weight[valid].sum()) <= 0:
            return None
        return float((score[valid] * weight[valid]).sum() / weight[valid].sum())
    score = score.dropna()
    return float(score.mean()) if len(score) else None


def _matching_observations(
    observations: list[dict], value: float | None, aggregation: str | None = None,
) -> list[dict]:
    if value is None:
        return []
    matches = []
    for observation in observations:
        if aggregation and observation["aggregation"] not in (
            aggregation, "unspecified",
        ):
            continue
        if abs(value - observation["value"]) <= observation["tolerance"]:
            matches.append(observation)
    return matches


def _choose_observation(matches: list[dict], score_column: str) -> dict | None:
    if not matches:
        return None
    score = _normalize_text(score_column)

    def rank(item):
        label = _normalize_text(item["label"])
        return (
            0 if label == score else 1,
            0 if "accuracy" in label or "точност" in label else 1,
            item["table_index"],
            item["row_index"],
        )

    return min(matches, key=rank)


def _majority_hypothesis(metric_df: pd.DataFrame, validation: dict) -> dict | None:
    """Доказанная отчётом формула majority vote, без agent-specific правил."""
    if not validation.get("majority_vote_cue"):
        return None
    scale_columns = _scale_columns(metric_df, _metric_columns(metric_df))
    sources = [
        column for column in scale_columns
        if column in validation.get("mentioned_columns", [])
        and not any(hint in _normalize_text(column)
                    for hint in _FINAL_METRIC_NAME_HINTS)
    ]
    if len(sources) < 3 or len(sources) > 7:
        return None
    derived = metric_df.copy()
    derived[DERIVED_SCORE_COLUMN] = _strict_majority(derived, sources)
    if derived[DERIVED_SCORE_COLUMN].notna().mean() < PARTIAL_MARKUP_MIN_FILL_SHARE:
        return None

    observations = validation.get("observations", [])
    macro = _measured_value(derived, DERIVED_SCORE_COLUMN)
    macro_matches = _matching_observations(observations, macro, "macro")
    weight_matches = []
    for column in _metric_columns(metric_df):
        if column in sources or column == DERIVED_SCORE_COLUMN:
            continue
        values = pd.to_numeric(metric_df[column], errors="coerce").dropna()
        weight_name = _normalize_text(column)
        has_weight_role = any(hint in weight_name for hint in (
            "weight", "freq", "frequency", "count", "вес", "частот",
            "количество", "число",
        ))
        if (values.empty or (values <= 0).any()
                or (values.nunique() <= MARKUP_SCALE_MAX_VALUES
                    and not has_weight_role)):
            continue
        weighted = _measured_value(derived, DERIVED_SCORE_COLUMN, column)
        for observation in _matching_observations(
            observations, weighted, "weighted"
        ):
            weight_matches.append((column, weighted, observation))

    unique_weights = {item[0] for item in weight_matches}
    if len(unique_weights) == 1:
        weight, value, observation = min(
            weight_matches, key=lambda item: (
                item[2]["table_index"], item[2]["row_index"]
            )
        )
        if macro_matches:
            return {
                "main_metric": DERIVED_SCORE_COLUMN,
                "metric_name": observation["label"],
                "other_metrics": [],
                "source_columns": sources,
                "row_aggregation": "majority_vote",
                "weight_column": weight,
                "strategy": "weighted_accuracy",
                "reported_validation_value": observation["value"],
                "resolution_source": "validation_report",
            }
    if macro_matches:
        observation = _choose_observation(macro_matches, DERIVED_SCORE_COLUMN)
        return {
            "main_metric": DERIVED_SCORE_COLUMN,
            "metric_name": observation["label"],
            "other_metrics": [],
            "source_columns": sources,
            "row_aggregation": "majority_vote",
            "weight_column": None,
            "strategy": "mean_score",
            "reported_validation_value": observation["value"],
            "resolution_source": "validation_report",
        }
    return None


def _direct_validation_hypothesis(
    metric_df: pd.DataFrame, validation: dict,
) -> dict | None:
    """Принять готовую score-колонку лишь после сверки с validation report."""
    if validation.get("majority_vote_cue"):
        return None
    candidates = _metric_columns(metric_df)
    score, _ = _resolve_by_data_properties(metric_df, candidates)
    if score is None and len(candidates) == 1:
        score = candidates[0]
    if score is None:
        return None
    weight = _detect_weight_by_product(metric_df, score)
    value = _measured_value(metric_df, score, weight)
    aggregation = "weighted" if weight else None
    matches = _matching_observations(
        validation.get("observations", []), value, aggregation
    )
    observation = _choose_observation(matches, score)
    if observation is None:
        return None
    other_metrics = []
    scale_columns = _scale_columns(metric_df, candidates)
    numeric = _numeric_by_name(metric_df)
    for item in validation.get("observations", []):
        column = _match_metric(item.get("label"), scale_columns)
        if not column or column == score or column in other_metrics:
            continue
        joint = pd.DataFrame({"main": numeric[score], "other": numeric[column]}).dropna()
        # Построчный дубль итоговой оценки не становится вторым критерием
        # ассессора лишь потому, что отчёт назвал его отдельно.
        if len(joint) and (joint["main"] == joint["other"]).all():
            continue
        other_metrics.append(column)
    return {
        "main_metric": score,
        "metric_name": observation["label"],
        "other_metrics": other_metrics,
        "source_columns": [score],
        "row_aggregation": "identity",
        "weight_column": weight,
        "strategy": "weighted_accuracy" if weight else "mean_score",
        "reported_validation_value": observation["value"],
        "resolution_source": "validation_report",
    }


def _materialize_metric_dataset(
    frame: pd.DataFrame, metric_df: pd.DataFrame, raw: dict, validation: dict,
) -> tuple[pd.DataFrame, pd.DataFrame, str | None]:
    aggregation = str(raw.get("row_aggregation") or "identity").strip()
    if aggregation == "identity":
        return frame, metric_df, None
    if aggregation != "majority_vote":
        return frame, metric_df, f"неподдержанная row_aggregation={aggregation!r}"
    if not validation.get("majority_vote_cue"):
        return frame, metric_df, (
            "majority_vote предложен без подтверждающей формулировки в "
            "отчёте о валидации"
        )
    sources = []
    for proposed in raw.get("source_columns") or []:
        matched = _match_metric(proposed, [str(column) for column in frame.columns])
        if matched and matched not in sources:
            sources.append(matched)
    if len(sources) < 3:
        return frame, metric_df, "для majority_vote нужны минимум три score-колонки"
    numeric = _numeric_by_name(metric_df)
    if any(source not in numeric or numeric[source].dropna().empty for source in sources):
        return frame, metric_df, "source_columns majority_vote не являются числовыми"
    if DERIVED_SCORE_COLUMN in frame.columns:
        return frame, metric_df, (
            f"зарезервированная колонка {DERIVED_SCORE_COLUMN!r} уже есть в корзине"
        )
    materialized = frame.copy()
    materialized[DERIVED_SCORE_COLUMN] = _strict_majority(materialized, sources)
    metric_materialized = metric_df.copy()
    metric_materialized[DERIVED_SCORE_COLUMN] = _strict_majority(
        metric_materialized, sources
    )
    raw["main_metric"] = DERIVED_SCORE_COLUMN
    raw["source_columns"] = sources
    raw["row_aggregation"] = "majority_vote"
    return materialized, metric_materialized, None


def _load_df(obj):
    """Принимает корзину в любом виде и возвращает DataFrame.

    Поддерживает: готовый DataFrame (как раньше); путь к .pkl/.pickle
    (распаковывается через pickle); путь к директории с таким файлом;
    сырые bytes пикла; .parquet/.xlsx/.csv по расширению.
    Это позволяет подавать на вход pickle, когда нет Data Source.
    """
    if isinstance(obj, pd.DataFrame):
        return obj
    if isinstance(obj, (bytes, bytearray)):
        return pickle.loads(obj)
    if isinstance(obj, (str, Path)):
        p = Path(obj)
        if p.is_dir():
            cand = sorted(
                [f for f in p.iterdir() if f.suffix.lower()
                 in (".pkl", ".pickle", ".parquet", ".xlsx", ".xls", ".csv")]
            )
            if not cand:
                raise FileNotFoundError(f"В директории '{p}' нет файла корзины.")
            p = cand[0]
        ext = p.suffix.lower()
        if ext in (".pkl", ".pickle"):
            return pd.read_pickle(str(p))
        if ext == ".parquet":
            return pd.read_parquet(str(p))
        if ext in (".xlsx", ".xls"):
            return pd.read_excel(str(p))
        if ext == ".csv":
            return pd.read_csv(str(p))
        # неизвестное расширение — пробуем pickle
        return pd.read_pickle(str(p))
    raise TypeError(f"Не могу загрузить корзину из объекта типа {type(obj)}")


def _instr_from_pickle_obj(data):
    """Pickle-объект инструкции → текст. dict с 'text' / строка / иное."""
    if isinstance(data, dict):
        return data.get("text", "") or data.get("instruction", "") or ""
    return str(data)


_NON_METRIC_COLS = {
    "query_id", "input_query", "output_answer", "question", "answer",
    "reference_answer", "scenario", "session_id", "history", "dialogue", "id",
    "trace_id", "span_id", "route", "starttime", "endtime", "start_time",
    "end_time", "distributive", "agent_id", "model_version_id",
    "input_query_count", "turn_index", "assessment_unit_id", "solution_version",
}


def _is_service_column(name):
    """Служебная колонка по имени (тексты корзины, комментарии, артефакты выгрузки).

    Заголовок, который сам читается как число, — не имя колонки, а значение,
    съехавшее в строку заголовков (сводный блок «метрика = 0.97…», записанный
    сбоку от таблицы разметки). Такая колонка не несёт построчных данных.
    """
    lc = str(name).strip().lower()
    if lc in _NON_METRIC_COLS or lc.startswith("unnamed"):
        return True
    if lc.endswith(("_id", "_comment", "_tool", "_text")):
        return True
    try:
        float(lc.replace(",", "."))
        return True
    except ValueError:
        return False


# Минимальная доля заполненных строк у колонки-кандидата в метрики. Построчная
# разметка покрывает корзину целиком или её размеченную часть (в корзинах
# пилота — не меньше половины строк: CI09877398 — 52%). Сводные агрегаты,
# дописанные сбоку таблицы, живут на строке-на-класс (CI09997438 — 4.5%
# строк). Порог 0.2 лежит между этими наблюдениями с запасом в обе стороны.
PARTIAL_MARKUP_MIN_FILL_SHARE = 0.2

# Категориальная разметка поддерживается самим assessor (точное совпадение
# меток), хотя среднее и динамика КМ для неё не вычислимы без отдельного
# числового отображения. Верхняя граница отсеивает свободный текст и ID; такой
# кандидат никогда не выбирается автоматически — только явной настройкой или
# точным ответом модели по фактическому имени колонки.
CATEGORICAL_MARKUP_MAX_VALUES = 20
_CATEGORICAL_METRIC_NAME_HINTS = (
    "оцен", "метрик", "критер", "итог", "результат", "вердикт",
    "соответств", "правильн", "релевант", "полнот", "точност", "безопас",
    "score", "metric", "label", "grade", "verdict", "quality", "result",
)


def _metric_columns(df):
    """Колонки-кандидаты в метрики корзины: все числовые непустые, кроме
    служебных. Суффикс *_metric/_score даёт приоритет в списке, но не
    отменяет проверку на числа и не вытесняет остальные критерии.

    Раньше одна колонка с суффиксом обрывала список — числовые
    бизнес-критерии без суффикса (safety, completeness) молча исчезали из
    ассессора и всех тестов.

    Доступ к колонкам позиционный: заголовок сводного блока в xlsx может
    прочитаться числом (float), и обращение df[str(имя)] к такой колонке
    падает KeyError ещё до всякой логики.
    """
    metric_like, numeric_plain = [], []
    total_rows = len(df) if hasattr(df, "__len__") else 0
    for position, raw_name in enumerate(getattr(df, "columns", [])):
        col = str(raw_name)
        if _is_service_column(col):
            continue
        series = df.iloc[:, position]
        values = pd.to_numeric(series.dropna(), errors="coerce")
        is_numeric = len(values) > 0 and values.notna().all()
        if is_numeric and total_rows:
            fill_share = len(values) / total_rows
            if fill_share < PARTIAL_MARKUP_MIN_FILL_SHARE:
                logging.warning(
                    "Колонка %r заполнена на %.0f%% строк — это сводный "
                    "агрегат или примечание, а не построчная разметка; "
                    "из кандидатов в метрики исключена", col, fill_share * 100,
                )
                continue
        if col.lower().endswith(("_metric", "_score", "_metrics")):
            if is_numeric:
                metric_like.append(col)
            else:
                logging.warning(
                    "Колонка %r названа как метрика, но не числовая — пропущена", col
                )
        elif is_numeric:
            numeric_plain.append(col)
    return metric_like + numeric_plain


def _categorical_metric_columns(df):
    """Категориальные колонки построчной разметки для assessor.

    Это не числовые кандидаты КМ: baseline и тест динамики по ним не считаются.
    Список нужен, чтобы корзина с метками вроде ``верно/неверно`` не падала до
    калибровки судьи. Служебные поля, свободный текст, константы и разреженные
    сводные блоки отсекаются по проверяемым свойствам данных.
    """
    candidates = []
    total_rows = len(df) if hasattr(df, "__len__") else 0
    for position, raw_name in enumerate(getattr(df, "columns", [])):
        col = str(raw_name)
        if _is_service_column(col):
            continue
        normalized_name = col.strip().lower().replace("ё", "е")
        if not any(hint in normalized_name for hint in _CATEGORICAL_METRIC_NAME_HINTS):
            # Низкая кардинальность сама по себе не означает разметку:
            # «Получатель», «канал», «тип продукта» и другие сегменты корзины
            # не должны попадать в prompt как кандидаты ключевой метрики.
            continue
        values = df.iloc[:, position].dropna()
        if values.empty or not total_rows:
            continue
        if len(values) / total_rows < PARTIAL_MARKUP_MIN_FILL_SHARE:
            continue
        # Числа (включая записанные строками) уже принадлежат _metric_columns.
        numeric = pd.to_numeric(values, errors="coerce")
        if numeric.notna().all():
            continue
        as_text = values.astype(str).str.strip()
        as_text = as_text[~as_text.str.lower().isin(("", "-", "—", "nan", "none", "null"))]
        if as_text.empty:
            continue
        unique = int(as_text.nunique())
        if 2 <= unique <= CATEGORICAL_MARKUP_MAX_VALUES:
            candidates.append(col)
    return candidates


def _match_metric(name, columns):
    """Сопоставить запрошенное имя метрики реальной колонке. None — если нет совпадения.
    Порядок: точное → регистр/пробелы → +/- суффикс _metric → подстрока."""
    if name is None:
        return None
    name = str(name).strip()
    by_lower = {c.lower().strip(): c for c in columns}
    if name in columns:
        return name
    if name.lower() in by_lower:
        return by_lower[name.lower()]
    # +/- суффикс _metric
    for cand in (name + "_metric", name.replace("_metric", ""), name + "_score"):
        if cand.lower() in by_lower:
            return by_lower[cand.lower()]
    # Подстрока (relevance ~ relevance_metric) — только для осмысленной длины
    # и только при ЕДИНСТВЕННОМ совпадении: раньше однобуквенная галлюцинация
    # LLM ('e') прилипала к первой попавшейся колонке.
    needle = name.lower()
    if len(needle) >= 4:
        hits = [orig for low, orig in by_lower.items()
                if needle in low or low in needle]
        if len(hits) == 1:
            return hits[0]
    return None


# Максимум различимых значений у шкалы построчной разметки. Шкалы корзин
# пилота: бинарная (2), 0–2 (3), 1–5 (5); 6 оставляет запас на шкалу 0–5.
# У весов, частот и производных произведений (оценка × частота) значений
# больше — колонкой построчной оценки они не бывают.
MARKUP_SCALE_MAX_VALUES = 6

# Минимум совместно размеченных строк, на которых проверяется тождество
# «итог = минимум частных критериев». На меньшей выборке равенство может
# оказаться совпадением, а не правилом разметки.
CONSOLIDATION_MIN_ROWS = 20

_FINAL_METRIC_NAME_HINTS = (
    "итог", "финаль", "общая оцен", "сводная оцен",
    "final", "overall", "total score", "total_score",
)


def _numeric_by_name(df):
    """Числовые значения колонок по строковому имени, доступ позиционный:
    имя-число (артефакт сводного блока) при df[str(имя)] даёт KeyError."""
    return {
        str(name): pd.to_numeric(df.iloc[:, position], errors="coerce")
        for position, name in enumerate(df.columns)
    }


def _scale_columns(df, candidates):
    """Кандидаты, чьи значения образуют шкалу разметки.

    Свойство данных, а не имени: немного различимых целых значений
    (MARKUP_SCALE_MAX_VALUES). Веса и частоты сценариев на реальных корзинах
    разнообразнее и отсекаются этим же свойством.
    """
    numeric = _numeric_by_name(df)
    result = []
    for col in candidates:
        values = numeric[col].dropna()
        if values.empty:
            continue
        if (values.nunique() <= MARKUP_SCALE_MAX_VALUES
                and (values == values.astype(int)).all()):
            result.append(col)
    return result


def _consolidated_column(df, scale_cols):
    """Колонка итоговой оценки среди шкальных кандидатов, если она одна.

    В корзинах пилота итог («Итог» CI09877398, final_mark CI09997554)
    считается как самый строгий из частных критериев: один проваленный
    критерий обнуляет итог. Это проверяемое тождество данных: колонка
    построчно равна минимуму остальных шкальных кандидатов везде, где
    размечены все. Побеждает только единственный кандидат — двусмысленность
    здесь означает дубли разметки, и решать её молча нельзя.
    """
    if len(scale_cols) < 3:
        return None
    numeric = _numeric_by_name(df)
    frame = pd.DataFrame({col: numeric[col] for col in scale_cols}).dropna()
    if len(frame) < CONSOLIDATION_MIN_ROWS:
        return None
    winners = []
    for col in scale_cols:
        others = [c for c in scale_cols if c != col]
        if (frame[col] == frame[others].min(axis=1)).all():
            winners.append(col)
    if not winners:
        return None
    if len(winners) == 1:
        return winners[0]
    # Несколько победителей возникает, когда частный критерий построчно
    # совпал с итогом (у CI09877398 «Релевантность» тождественна «Итогу»:
    # при проваленной классификации релевантность тоже 0). Тождественные
    # значения — одно и то же измерение, записанное дважды; итоговую
    # колонку дописывают правее частных критериев, берётся правая.
    # Различающиеся победители — настоящая двусмысленность, её решает
    # только явная настройка.
    reference = frame[winners[0]]
    if all((frame[col] == reference).all() for col in winners[1:]):
        return winners[-1]
    return None


def _detect_weight_by_product(df, main):
    """Колонка веса, доказанная самой корзиной, если она единственная.

    Команды агентов считают взвешенную метрику прямо в файле разметки:
    рядом с оценкой лежит колонка-произведение (оценка × частота сценария),
    как «mark * freq» у CI09997438. Если в корзине есть колонка, построчно
    равная произведению ключевой метрики на другую числовую колонку, —
    взвешивание задумано автором разметки, и вес восстанавливается без
    участия модели. Неоднозначность (несколько кандидатов веса) означает,
    что доказательства нет.
    """
    numeric = _numeric_by_name(df)
    if main not in numeric:
        return None
    candidates = [col for col in _metric_columns(df) if col != main]
    main_values = numeric[main]
    weights = set()
    for product_col in candidates:
        for weight_col in candidates:
            if weight_col == product_col:
                continue
            joint = pd.DataFrame({
                "product": numeric[product_col],
                "weight": numeric[weight_col],
                "main": main_values,
            }).dropna()
            if len(joint) < CONSOLIDATION_MIN_ROWS:
                continue
            if (joint["weight"] > 0).all() and (
                    joint["product"] == joint["main"] * joint["weight"]).all():
                weights.add(weight_col)
    return weights.pop() if len(weights) == 1 else None


def _resolve_by_data_properties(df, candidates):
    """Детерминированный выбор метрики, когда ответ модели не совпал с колонками.

    Порядок: единственная явно итоговая шкальная колонка → единственная
    консолидированная итоговая оценка → единственная шкальная колонка
    разметки. Имя используется только как однозначный финальный маркер среди
    уже доказанных шкальных колонок, а не как словарь произвольных метрик.
    None — свойства данных однозначного ответа не дают.
    """
    scale_cols = _scale_columns(df, candidates)
    final_named = [
        column for column in scale_cols
        if any(
            hint in column.strip().lower().replace("ё", "е")
            for hint in _FINAL_METRIC_NAME_HINTS
        )
    ]
    if len(final_named) == 1:
        return final_named[0], (
            f"единственная шкальная колонка с явным признаком итоговой "
            f"оценки в имени: {final_named[0]!r}"
        )
    consolidated = _consolidated_column(df, scale_cols)
    if consolidated:
        return consolidated, (
            f"колонка {consolidated!r} построчно равна минимуму остальных "
            f"шкальных критериев {sorted(set(scale_cols) - {consolidated})} — "
            "это итоговая оценка разметки"
        )
    if len(scale_cols) == 1:
        return scale_cols[0], (
            f"единственная колонка со шкалой разметки среди кандидатов "
            f"{candidates}: {scale_cols[0]!r}"
        )
    return None, None


def _parse_metrics_json(text):
    """Достать dict с критериями из ответа модели, даже если он обёрнут в
    ```json ... ``` или вокруг есть лишний текст."""
    import re as _re
    if isinstance(text, dict):
        return text
    s = str(text).strip()
    if "```" in s:
        m = _re.search(r"```(?:json)?\s*(.*?)```", s, _re.S)
        if m:
            s = m.group(1).strip()
    try:
        return json.loads(s)
    except Exception:
        pass
    m = _re.search(r"\{.*\}", s, _re.S)   # первый сбалансированный объект
    if m:
        return json.loads(m.group(0))
    raise ValueError("в ответе модели не найден JSON")


def _chat_content(payload):
    """Текст OpenAI-compatible ответа, включая список content-блоков.

    MiniMax при исчерпанном лимите рассуждения может вернуть пустой content.
    Такой ответ не маскируется под ошибку JSON: в исключение попадают
    finish_reason и request id, по которым сбой виден в логе AI Gateway.
    """
    try:
        choice = payload["choices"][0]
        message = choice["message"]
    except (KeyError, IndexError, TypeError) as exc:
        raise ValueError("AI Gateway вернул ответ без choices[0].message") from exc
    content = message.get("content") if isinstance(message, dict) else None
    if isinstance(content, str):
        text = content
    elif isinstance(content, list):
        text = "".join(
            block if isinstance(block, str) else block.get("text", "")
            for block in content
            if isinstance(block, (str, dict))
        )
    else:
        text = ""
    if text.strip():
        return text
    raise ValueError(
        "AI Gateway вернул пустой message.content "
        f"(finish_reason={choice.get('finish_reason')!r}, "
        f"request_id={payload.get('id')!r})"
    )


def _walk_dicts(value):
    """Все словари вложенного результата doc-browser, без привязки к обёртке."""
    if isinstance(value, dict):
        yield value
        for nested in value.values():
            yield from _walk_dicts(nested)
    elif isinstance(value, (list, tuple)):
        for nested in value:
            yield from _walk_dicts(nested)


def _parse_optional_float(value):
    """Число из структурированного поля doc-browser (`0.8`, `80%`, `0,8`)."""
    if value is None or isinstance(value, bool):
        return None
    text = str(value).strip().replace(" ", "").replace(",", ".")
    if not text or text == "-":
        return None
    percent = text.endswith("%")
    if percent:
        text = text[:-1]
    try:
        result = float(text)
    except ValueError:
        return None
    return result / 100.0 if percent else result


def _doc_browser_context(value):
    """Карточка doc-browser → текст для модели и проверенные поля метрики.

    Актуальный g-aiva-doc-browser уже извлекает `evaluation_metric`,
    `threshold`, `metric_desc`, `metric_func` в
    `extracted_fields.bp_card_fields`. Повторно просить LLM прочитать эти же
    поля не только дорого, но и менее надёжно. Функция понимает и новый
    структурированный порт, и прежний `all_results.bp_card`.
    """
    if value is None:
        return "", {}
    if isinstance(value, str):
        return value, {}

    fields = {}
    text = ""
    if isinstance(value, dict):
        # Сначала специально ищем канонический блок; затем принимаем плоскую
        # форму, если платформа сняла значение прямо с extracted_fields.
        candidates = []
        for mapping in _walk_dicts(value):
            nested = mapping.get("bp_card_fields")
            if isinstance(nested, dict):
                candidates.insert(0, nested)
            # Одинокий ключ threshold во вложенных extraction-полях может
            # относиться к алгоритму/гиперпараметру, а не к бизнес-метрике.
            # Кандидат обязан нести семантику самой метрики.
            if any(key in mapping for key in (
                "evaluation_metric", "metric_name", "metric_desc", "metric_func"
            )):
                candidates.append(mapping)
            for key in ("bp_card", "card", "text"):
                candidate_text = mapping.get(key)
                if not text and isinstance(candidate_text, str) and candidate_text.strip():
                    text = candidate_text

        for candidate in candidates:
            for key in ("evaluation_metric", "metric_desc", "metric_func"):
                raw = candidate.get(key)
                if key not in fields and raw not in (None, "", "-"):
                    fields[key] = str(raw).strip()
            threshold = _parse_optional_float(candidate.get("threshold"))
            if "threshold" not in fields and threshold is not None:
                fields["threshold"] = threshold

    if fields:
        structured = json.dumps(fields, ensure_ascii=False)
        text = (text + "\n\nСтруктурированные поля отчёта: " + structured).strip()
    return text, fields


def _metric_contract(raw, resolved, df, *, explicit_metric, document_fields):
    """Расширить legacy-ответ selector до единого исполнимого MetricSpec.

    Каждая колонка повторно сопоставляется с реальной корзиной. Стратегия
    выводится из проверенных ролей, поэтому галлюцинация модели вроде
    `weighted_accuracy` без колонки веса наружу не проходит.
    """
    all_columns = [str(column) for column in df.columns]
    numeric_columns = _metric_columns(df)
    categorical_columns = _categorical_metric_columns(df)
    main = resolved["main_metric"]
    is_categorical = main in categorical_columns and main not in numeric_columns

    def numeric_role(*names):
        for name in names:
            matched = _match_metric(raw.get(name), numeric_columns)
            if matched and matched != main:
                return matched
        return None

    def any_role(*names):
        for name in names:
            matched = _match_metric(raw.get(name), all_columns)
            if matched:
                return matched
        return None

    weight = numeric_role("weight_column", "weight_col")
    if weight is None:
        weight = _detect_weight_by_product(df, main)
        if weight:
            logging.info(
                "Колонка веса %r восстановлена по колонке-произведению в "
                "корзине (оценка × вес)", weight,
            )
    prediction = any_role("prediction_column", "prediction_col")
    reference = any_role("reference_column", "reference_col")
    proposed_strategy = str(raw.get("strategy") or "").strip()
    if is_categorical:
        # Assessor умеет калиброваться и размечать номинальные классы. Среднее
        # такой шкалы не определено, поэтому зависимые числовые тесты увидят
        # отдельный статус и честно вернут not_computable.
        weight = None
        strategy = "categorical_label"
    elif weight:
        strategy = "weighted_accuracy"
    elif proposed_strategy == "accuracy" and prediction and reference:
        strategy = "accuracy"
    else:
        # Selector всегда резолвит main_metric в числовую score-колонку.
        # При материализованной построчной оценке её baseline — среднее.
        strategy = "mean_score"

    metric_name = (
        document_fields.get("evaluation_metric")
        or raw.get("metric_name")
        or raw.get("evaluation_metric")
        or main
    )
    source = str(
        raw.get("resolution_source")
        or resolved.get("resolution_source", "llm")
    )
    row_aggregation = str(raw.get("row_aggregation") or "identity").strip()
    if row_aggregation not in ROW_AGGREGATIONS:
        row_aggregation = "identity"
    source_columns = []
    for proposed in raw.get("source_columns") or [main]:
        matched = _match_metric(proposed, all_columns)
        if matched and matched not in source_columns:
            source_columns.append(matched)
    if not source_columns:
        source_columns = [main]

    scale_min = _parse_optional_float(raw.get("scale_min"))
    scale_max = _parse_optional_float(raw.get("scale_max"))
    warnings = []
    if proposed_strategy and proposed_strategy not in METRIC_STRATEGIES:
        warnings.append(
            f"предложенная моделью стратегия {proposed_strategy!r} не поддержана; "
            f"использована проверяемая стратегия {strategy!r}"
        )

    raw_direction = raw.get("higher_is_better", True)
    if isinstance(raw_direction, bool):
        higher_is_better = raw_direction
    else:
        higher_is_better = str(raw_direction).strip().lower() not in (
            "false", "0", "no", "нет",
        )

    return {
        "metric_spec_version": METRIC_SPEC_VERSION,
        "status": "resolved_categorical" if is_categorical else "resolved",
        # Legacy-поля для существующих потребителей.
        "main_metric": main,
        "other_metrics": resolved["other_metrics"],
        # Исполнимый контракт для baseline-extractor.
        "metric_name": str(metric_name).strip(),
        "score_column": main,
        "source_columns": source_columns,
        "row_aggregation": row_aggregation,
        "marker_columns": resolved["other_metrics"],
        "weight_column": weight,
        "prediction_column": prediction,
        "reference_column": reference,
        "strategy": strategy,
        "score_value_type": "categorical" if is_categorical else "numeric",
        "scale_min": scale_min,
        "scale_max": scale_max,
        "higher_is_better": higher_is_better,
        "business_threshold": document_fields.get("threshold"),
        "metric_description": document_fields.get("metric_desc"),
        "metric_formula": document_fields.get("metric_func"),
        "resolution_source": source,
        "reported_validation_value": _parse_optional_float(
            raw.get("reported_validation_value")
        ),
        "evidence": [
            f"колонка {main!r} сопоставлена с фактической схемой корзины",
            f"стратегия {strategy!r} выведена из проверенных ролей колонок",
        ],
        "warnings": warnings,
    }


def _not_computable_metric_contract(raw, resolved, df, *, document_fields):
    """Нефатальный MetricSpec для корзины без однозначной числовой КМ.

    Legacy-потребители требуют непустой ``main_metric``. Канонический
    ``target`` здесь является только техническим именем пустой колонки,
    которую assessor добавит к мониторинговым строкам; ``score_column=None``
    и ``status=not_computable`` не позволяют принять её за измерение.
    """
    metric_name = (
        document_fields.get("evaluation_metric")
        or raw.get("metric_name")
        or raw.get("evaluation_metric")
        or "Ключевая метрика"
    )
    reason = resolved.get("reason") or (
        "корзина не содержит однозначной построчной колонки ключевой метрики"
    )
    numeric = _metric_columns(df)
    categorical = _categorical_metric_columns(df)
    return {
        "metric_spec_version": METRIC_SPEC_VERSION,
        "status": "not_computable",
        # Транспортный placeholder: downstream сохраняет строки и формирует
        # серые результаты; фактической score-колонкой он не объявлен.
        "main_metric": "target",
        "other_metrics": [],
        "metric_name": str(metric_name).strip(),
        "score_column": None,
        "source_columns": [],
        "row_aggregation": None,
        "marker_columns": [],
        "weight_column": None,
        "prediction_column": None,
        "reference_column": None,
        "strategy": None,
        "score_value_type": "unknown",
        "scale_min": None,
        "scale_max": None,
        "higher_is_better": True,
        "business_threshold": document_fields.get("threshold"),
        "metric_description": document_fields.get("metric_desc"),
        "metric_formula": document_fields.get("metric_func"),
        "resolution_source": "not_computable",
        "reported_validation_value": None,
        "reason_code": resolved.get("reason_code", "metric_column_unresolved"),
        "reason": reason,
        "evidence": [
            f"колонки корзины: {[str(column) for column in df.columns]}",
            f"числовые кандидаты: {numeric}",
            f"категориальные кандидаты: {categorical}",
        ],
        "warnings": [
            reason,
            "числовые тесты качества должны вернуть серый not_computable; "
            "техническая колонка target не является измеренной метрикой",
        ],
    }


def _resolve_main_metric(json_output, df, explicit_metric=""):
    """Определяет колонку ключевой метрики. Порядок:

    1. явная настройка узла (`main_metric` в Settings) — приоритет всегда за ней;
    2. ответ LLM, совпавший с реальной колонкой корзины;
    3. единственный кандидат среди колонок-метрик;
    4. свойства данных: консолидированная итоговая оценка (построчный минимум
       частных критериев) либо единственная шкальная колонка разметки;
    5. точное совпадение с категориальной разметкой разрешает только assessor;
    6. настоящая неоднозначность возвращает ``not_computable`` без остановки.

    Модель отвечает семантическим именем метрики («взвешенная accuracy»), а не
    колонкой, регулярно — падать из-за этого нельзя ни в контуре, ни локально.
    Пайплайн при этом не угадывает: если и свойства данных однозначного ответа
    не дают — узел возвращает машинный отказ. Это критично для
    невозобновляемой автовалидации: workflow доходит до отчёта, но не получает
    ложный зелёный результат.
    """
    cols = _metric_columns(df)
    categorical_cols = _categorical_metric_columns(df)
    source = "llm"

    if explicit_metric:
        main = _match_metric(explicit_metric, [str(c) for c in df.columns])
        source = "setting"
        if main is None:
            return {
                "main_metric": None,
                "other_metrics": [],
                "resolution_source": "not_computable",
                "reason_code": "explicit_metric_missing",
                "reason": (
                    f"колонка метрики {explicit_metric!r} из Settings не найдена "
                    f"в корзине; колонки: {list(df.columns)}"
                ),
            }
    else:
        main = _match_metric(json_output.get("main_metric"), cols)
        if main is None:
            categorical = _match_metric(
                json_output.get("main_metric"), categorical_cols
            )
            if categorical is not None:
                main = categorical
                source = "llm_categorical"
        if main is None and len(cols) == 1:
            main = cols[0]
            source = "single_numeric_candidate"
            logging.warning(
                "Ответ модели %r не совпал с колонками; взят единственный "
                "кандидат %r", json_output.get("main_metric"), main,
            )
        if main is None:
            main, reason = _resolve_by_data_properties(df, cols)
            if main is not None:
                source = "data_properties"
                logging.warning(
                    "Ответ модели %r не совпал с колонками; метрика %r "
                    "определена по свойствам данных: %s",
                    json_output.get("main_metric"), main, reason,
                )
        if main is None:
            if cols:
                hint = f"числовых кандидатов несколько: {cols}"
                reason_code = "ambiguous_numeric_candidates"
            elif categorical_cols:
                hint = (
                    "числовых кандидатов нет; категориальные колонки "
                    f"{categorical_cols} не совпали с ответом модели"
                )
                reason_code = "categorical_metric_unresolved"
            else:
                hint = "в корзине нет построчных колонок разметки"
                reason_code = "no_markup_candidates"
            return {
                "main_metric": None,
                "other_metrics": [],
                "resolution_source": "not_computable",
                "reason_code": reason_code,
                "reason": (
                    "не удалось однозначно определить колонку ключевой метрики: "
                    f"ответ модели {json_output.get('main_metric')!r} не совпал "
                    f"с колонками корзины, {hint}. Для числового мониторинга "
                    "нужна размеченная эталонная корзина либо точное main_metric"
                ),
            }

    # Доп. критерии матчим по ВСЕМ колонкам корзины: сжатый список кандидатов
    # раньше делал упомянутый моделью критерий невосстановимым. Невматченные
    # имена логируются, а не пропадают молча.
    others = []
    all_columns = [str(c) for c in df.columns]
    for m in json_output.get("other_metrics", []):
        mm = _match_metric(m, all_columns)
        if mm and mm != main and mm not in others:
            others.append(mm)
        elif not mm:
            logging.warning("Критерий %r из ответа модели не найден в корзине", m)
    # Наружу — только резолвленные имена: chain-of-thought модели не должен
    # утекать в порт и дальше во все артефакты.
    return {"main_metric": main, "other_metrics": others,
            "resolution_source": source}


def _load_instruction(obj):
    """Инструкция по разметке в любом виде → текст.

    КРИТИЧНО: тип определяется по СИГНАТУРЕ БАЙТОВ файла, а не по расширению,
    потому что платформенный getPortAsLocalPath часто отдаёт файл под именем без
    расширения (или с чужим), и тогда .pickle ошибочно читается как .docx.
    Сигнатуры: ZIP/DOCX = b'PK\\x03\\x04'; pickle (proto 2-5) = b'\\x80'.
    Поддерживает: dict/строку; bytes; путь к .docx; путь к .pkl/.pickle;
    путь без расширения (определяется по содержимому); директорию с любым из них.
    """
    import pickle as _pickle
    if obj is None:
        return ""
    if isinstance(obj, dict):
        return _instr_from_pickle_obj(obj)
    if isinstance(obj, (bytes, bytearray)):
        b = bytes(obj)
        if b[:4] == b"PK\x03\x04":           # это zip/docx в байтах
            # Времянка с уникальным именем: фиксированный путь в /tmp при двух
            # узлах на одной машине приводил к чтению чужой инструкции.
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(b)
            try:
                return read_docx(tmp.name)
            finally:
                os.unlink(tmp.name)
        return _instr_from_pickle_obj(_pickle.loads(b))  # иначе pickle

    p = Path(obj)
    # директория: выбрать единственный/подходящий файл
    if p.is_dir():
        files = [f for f in p.iterdir() if f.is_file() and not f.name.startswith((".", "~$"))]
        pkls = [f for f in files if f.suffix.lower() in (".pkl", ".pickle")]
        docs = [f for f in files if f.suffix.lower() == ".docx"]
        if pkls:
            p = sorted(pkls)[0]
        elif docs:
            return read_docx(str(p))
        elif len(files) == 1:
            p = files[0]               # одно имя без расширения — определим по сигнатуре
        else:
            return read_docx(str(p))   # пусть read_docx сам разберётся/сообщит

    if not p.is_file():
        raise FileNotFoundError(f"Инструкция: путь не найден: {p}")

    # ОПРЕДЕЛЕНИЕ ПО СИГНАТУРЕ (расширение может быть потеряно платформой)
    with open(str(p), "rb") as f:
        head = f.read(4)
    if not head:
        return ""   # пустой/0-байтный файл инструкции — инструкция опциональна
    if head[:4] == b"PK\x03\x04":              # zip → docx
        return read_docx(str(p))
    if head[:1] == b"\x80" or p.suffix.lower() in (".pkl", ".pickle"):  # pickle
        with open(str(p), "rb") as f:
            return _instr_from_pickle_obj(_pickle.load(f))
    if p.suffix.lower() == ".docx":
        return read_docx(str(p))
    # последний шанс: сначала pickle, потом docx (с понятной ошибкой)
    try:
        with open(str(p), "rb") as f:
            return _instr_from_pickle_obj(_pickle.load(f))
    except Exception:
        return read_docx(str(p))


def read_docx(path):
    """Читает один или несколько DOCX-файлов без записи во входную директорию."""
    path = Path(path)
    if path.is_file():
        files = [path]
    elif path.is_dir():
        files = sorted(
            f for f in path.iterdir()
            if f.is_file()
            and f.suffix.lower() == ".docx"
            and not f.name.startswith("~$")
            and not f.name.startswith(".")
        )
    else:
        raise FileNotFoundError(f"Путь с docx-инструкцией не найден: {path}")

    if not files:
        raise FileNotFoundError(f"В '{path}' нет DOCX-файлов для чтения.")

    texts = []
    for source_file in files:
        try:
            doc = Document(str(source_file))
            text = "\n".join(para.text for para in doc.paragraphs)
            texts.append(text)
        except Exception as e:
            raise RuntimeError(f"Ошибка при чтении DOCX '{source_file}': {e}") from e

    return "\n\n".join(texts)


def make_giga_request(input_prompt, model: str = ""):
    """Запрос к GigaChat; имя модели приходит из настройки узла."""
    from langchain_gigachat.chat_models import GigaChat

    from config import ModelsConfig

    config = ModelsConfig(model=model)
    chat = GigaChat(**config.contour_llm_configs)

    return chat.invoke(input_prompt)


def make_sds_request(model_id, input_prompt, system_prompt):
    model = model_id
    gateway_url = os.environ.get("AI_GATEWAY_URL", None)
    if not gateway_url:
        raise ValueError("AI_GATEWAY_URL is required for SDS model_id != 'giga'")
    chat_url = f"{gateway_url.rstrip('/')}/api/v1/chat/completions"

    common_headers = {"Content-Type": "application/json"}
    DEFAULT_COT_SYSTEM_PROMPT = (
        "Ты — умный ассистент. Всегда рассуждай шаг за шагом, объясняя свою мысль."
    )

    def generate_chat_payload(
        messages: list,
        temperature: float = 1.0,
        top_p: float = 0.1,
        top_k: int = -1,  # -1 = отключено
        presence_penalty: float = 0.0,
        frequency_penalty: float = 0.0,
        max_tokens: int = 512,
        repetition_penalty: float = 1.0,
        cot_enabled: bool = False,  # Включение chain-of-thought
        long_context: bool = False,  # Активация 128k контекста
        structured_output=None,  # None, True, или schema dict
    ):
        # Клонируем сообщения, чтобы не менять оригинал

        final_messages = [msg.copy() for msg in messages]

        # Добавляем системный промпт для CoT
        if cot_enabled:
            has_system = any(msg["role"] == "system" for msg in final_messages)

            if not has_system:
                final_messages.insert(
                    0, {"role": "system", "content": DEFAULT_COT_SYSTEM_PROMPT}
                )

        # Формируем базовый запрос

        payload = {
            "model": model,
            "messages": final_messages,
            "temperature": temperature,
            "top_p": top_p,
            "n": 1,
            "stream": False,
            "max_tokens": max_tokens,
            "repetition_penalty": repetition_penalty,
            "top_k": top_k,
            "presence_penalty": presence_penalty,
            "frequency_penalty": frequency_penalty,
        }

        # Параметр для длинного контекста (если поддерживается vLLM)

        if long_context:
            payload["max_model_len"] = 131072  # 128k токенов

        # Structured output, если запрошен

        if structured_output is not None:
            payload["structured_output"] = structured_output

        return json.dumps(payload)

    def generate(system_message: str, user_question: str, json_schema: str) -> str:
        chat_payload = generate_chat_payload(
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_question},
            ],
            temperature=0.1,
            top_p=0.1,
            top_k=1,
            presence_penalty=0.5,
            frequency_penalty=0.5,
            max_tokens=8192,
            cot_enabled=True,
            long_context=True,
            structured_output=json_schema,
        )

        # Выберите какую переменную отправлять ниже (chat_payload или chat_payload_schema)

        logging.info(
            "kriteria-selector AI Gateway request: model=%s prompt_chars=%s "
            "max_tokens=%s timeout=%ss",
            model, len(system_message) + len(user_question), 8192, 300,
        )
        started = time.monotonic()
        chat_response = requests.post(
            chat_url,
            headers=common_headers,
            data=chat_payload,  # или chat_payload_schema
            verify=False,
            timeout=(10, 300),
        )
        chat_response.raise_for_status()
        answer = json.loads(chat_response.text)
        try:
            choice = answer["choices"][0]
            content = choice.get("message", {}).get("content")
            content_chars = len(content) if isinstance(content, str) else 0
            usage = answer.get("usage") or {}
            logging.info(
                "kriteria-selector AI Gateway response: model=%s http=%s "
                "elapsed=%.2fs request_id=%s finish_reason=%s "
                "completion_tokens=%s content_chars=%s",
                model, chat_response.status_code, time.monotonic() - started,
                answer.get("id"), choice.get("finish_reason"),
                usage.get("completion_tokens"), content_chars,
            )
        except (KeyError, IndexError, TypeError):
            logging.warning(
                "kriteria-selector AI Gateway response has unexpected shape: "
                "model=%s http=%s elapsed=%.2fs",
                model, chat_response.status_code, time.monotonic() - started,
            )
        return answer

    return generate(
        system_message=system_prompt,
        user_question=input_prompt,
        json_schema=None,
    )


def _truncate_cells(frame: pd.DataFrame) -> pd.DataFrame:
    """Пример данных для промта: длинные значения обрезаются.

    Диалоговые корзины несут реплики на тысячи символов; в промте они
    вытесняют инструкцию разметки, при том что для определения роли колонки
    хватает начала значения.
    """
    def _shorten(value):
        text = str(value)
        if len(text) <= SAMPLE_VALUE_MAX_CHARS:
            return value
        return text[:SAMPLE_VALUE_MAX_CHARS] + "…"

    return frame.map(_shorten) if hasattr(frame, "map") else frame


def _metric_rows(frame: pd.DataFrame) -> pd.DataFrame:
    """Строки реальных примеров, на которых определена построчная метрика.

    В Excel рядом с таблицей встречаются сводные baseline/частоты без вопроса
    и ответа. Они полезны отдельному baseline-extractor, но не являются новым
    классом разметки и не должны мешать детерминированно выбрать итоговую
    колонку. Обычные строки с частичными метками сохраняются.
    """
    for question, answer in (
        ("input_query", "output_answer"),
        ("question", "answer"),
    ):
        if question not in frame.columns or answer not in frame.columns:
            continue
        mask = pd.Series(True, index=frame.index)
        for column in (question, answer):
            mask &= frame[column].notna()
            mask &= frame[column].map(lambda value: bool(str(value).strip()))
        if mask.any():
            removed = int((~mask).sum())
            if removed:
                logging.info(
                    "kriteria-selector: %s сводных/пустых строк без Q/A не "
                    "участвуют в выборе построчной метрики", removed,
                )
            return frame[mask].copy()
    return frame


def _finalize_metric_selection(
    json_output,
    metric_df,
    *,
    explicit_metric,
    document_fields,
    artifact_provenance,
    validation_evidence,
    metric_dataset,
    model_id,
    llm_called=False,
):
    """Собрать единый MetricSpec после выбора источника ключевой метрики."""
    resolved = _resolve_main_metric(
        json_output, metric_df, explicit_metric=explicit_metric
    )
    if resolved.get("main_metric") is None:
        result = _not_computable_metric_contract(
            json_output, resolved, metric_df, document_fields=document_fields,
        )
    else:
        result = _metric_contract(
            json_output,
            resolved,
            metric_df,
            explicit_metric=explicit_metric,
            document_fields=document_fields,
        )
    validation = {
        key: value for key, value in validation_evidence.items()
        if key != "text_for_prompt"
    }
    if result["status"] == "resolved" and validation.get("provided"):
        if validation.get("status") == "unreadable":
            raise ArtifactTransportError(
                "validation_report_artifact помечен как unreadable: "
                f"{validation.get('reason', 'неизвестная ошибка')}"
            )
        elif validation.get("observations"):
            aggregation = "weighted" if result.get("weight_column") else None
            measured = _measured_value(
                metric_df,
                result["score_column"],
                result.get("weight_column"),
            )
            matches = _matching_observations(
                validation["observations"], measured, aggregation
            )
            selected = _choose_observation(matches, result["score_column"])
            if selected is None:
                reason = (
                    "формула MetricSpec не воспроизводит ни одного значения "
                    "ключевой метрики из отчёта о валидации"
                )
                validation.update({
                    "status": "contradiction",
                    "computed_value": measured,
                })
                result.update({
                    "status": "not_computable",
                    "main_metric": "target",
                    "score_column": None,
                    "strategy": None,
                    "reason_code": "validation_metric_not_reproduced",
                    "reason": reason,
                })
                result["warnings"].append(reason)
            else:
                validation.update({
                    "status": "confirmed",
                    "computed_value": measured,
                    "selected_observation": selected,
                })
                result["reported_validation_value"] = selected["value"]
                result["evidence"].append(
                    f"пересчёт {measured:.12g} воспроизводит опубликованное "
                    f"значение {selected['value']:g} ({selected['label']!r}) "
                    f"с допуском {selected['tolerance']:g}"
                )
        else:
            validation["status"] = "no_metric_observation"
            result["warnings"].append(
                "отчёт о валидации прочитан, но типизированная таблица "
                "ключевой метрики со столбцом «Значение» не найдена"
            )
    result["validation_evidence"] = validation
    result["artifact_provenance"] = artifact_provenance
    result["selector_inference"] = {
        "llm_called": bool(llm_called),
        "configured_model_id": model_id,
        "selection_path": result["resolution_source"],
    }
    logging.info(
        "kriteria-selector MetricSpec: status=%s metric_name=%r "
        "score_column=%r strategy=%s source=%s llm_called=%s model=%s",
        result["status"], result["metric_name"], result["score_column"],
        result["strategy"], result["resolution_source"], llm_called, model_id,
    )
    return {
        "metric_spec": result,
        "metric_dataset": metric_dataset,
    }


def _load_monitoring_contract(value) -> dict | None:
    """Восстановить monitoring_metric из транспорта SberDS.

    Порт ``default`` может вернуть готовый dict, JSON/pickle bytes либо путь
    без расширения на скачанный файл. Подключённый, но нераспознанный контракт
    не должен незаметно отправлять selector в legacy LLM-ветку.
    """
    if value is None:
        return None
    current = value
    for _ in range(_ARTIFACT_MAX_UNWRAP_DEPTH):
        if isinstance(current, dict):
            # Контракт not_computable не несёт scoring — это законный отказ
            # адаптера, а не нераспознанный транспорт.
            if isinstance(current.get("scoring"), dict) or (
                current.get("contract_version") == "laim-monitoring-metric.v2"
                and current.get("status")
            ):
                return current
            nested = next(
                (
                    current[key]
                    for key in (
                        "monitoring_metric", "result", "payload", "data", "value"
                    )
                    if key in current and current[key] is not current
                ),
                None,
            )
            if nested is None:
                break
            current = nested
            continue
        if isinstance(current, (bytes, bytearray)):
            raw = bytes(current)
            try:
                current = json.loads(raw.decode("utf-8-sig"))
                continue
            except (UnicodeDecodeError, json.JSONDecodeError):
                try:
                    current = pickle.loads(raw)
                    continue
                except Exception:
                    break
        if isinstance(current, (str, Path)):
            text = str(current).strip()
            if text.startswith("{"):
                try:
                    current = json.loads(text)
                    continue
                except json.JSONDecodeError:
                    break
            path = Path(text)
            if path.is_dir():
                files = sorted(item for item in path.iterdir() if item.is_file())
                if len(files) != 1:
                    break
                path = files[0]
            if path.is_file():
                current = path.read_bytes()
                continue
        break
    raise ValueError(
        "monitoring_metric подключён, но контракт laim-monitoring-metric.v2 "
        "не прочитан; ожидается dict, JSON/pickle bytes или локальный файл"
    )


def _selector_from_monitoring_contract(metric, df) -> dict | None:
    """MetricSpec напрямую из laim-monitoring-metric.v2 без LLM и артефактов.

    Метрика уже определена MeasurementPlan-ом laim-baskets-adapter со сверкой
    baseline против validation report — селектору достаточно передать её
    дальше. Возврат None означает «контракт не передан/не распознан» и
    отправляет узел в прежний legacy-путь (LLM + отчёты).
    """
    if not isinstance(metric, dict):
        return None
    if metric.get("contract_version") != "laim-monitoring-metric.v2" and "scoring" not in metric:
        return None
    status = str(metric.get("status") or "").strip().lower()
    if status != "computed":
        # Отказ адаптера передаётся как отказ: судить итоговую колонку без
        # объявленной КМ нельзя, конвертер ниже всё равно остановится.
        return {
            "status": "not_computable",
            "main_metric": None,
            "other_metrics": [],
            "metric_name": metric.get("name"),
            "score_column": None,
            "source_columns": [],
            "marker_columns": [],
            "row_aggregation": None,
            "strategy": "monitoring_metric_passthrough",
            "scoring_method": ((metric.get("scoring") or {}).get("method")),
            "assessment_mode": metric.get("assessment_mode"),
            "weight_column": None,
            "reported_validation_value": None,
            "resolution_source": "monitoring_metric_not_computable",
            "reason_code": metric.get("reason_code"),
            "reason": metric.get("reason") or f"контракт monitoring_metric со статусом {status!r}",
            "warnings": [],
        }
    scoring = metric.get("scoring") or {}
    sources = scoring.get("sources") or []
    columns, roles = [], {}
    for source in sources:
        name = str(source.get("column_name", "")).strip()
        if name and name not in columns:
            columns.append(name)
            roles[name] = str(source.get("role", "")).strip()
    method = str(scoring.get("method") or "identity").strip().lower()
    baseline = metric.get("baseline") or {}

    def _spec(main, others, source_path, reason=None):
        all_columns = list(dict.fromkeys(
            name for name in [main, *others] if name
        ))
        aggregate_sources = method in {"majority", "mean_criteria", "all_criteria"}
        spec = {
            "status": "resolved",
            "main_metric": main,
            "other_metrics": [name for name in others if name != main],
            "metric_name": metric.get("name"),
            "score_column": main,
            "source_columns": all_columns if aggregate_sources else [main],
            "marker_columns": [name for name in others if name != main],
            "row_aggregation": "majority_vote" if method == "majority" else "identity",
            "strategy": "monitoring_metric_passthrough",
            "scoring_method": method,
            "assessment_mode": metric.get("assessment_mode"),
            "weight_column": (metric.get("aggregation") or {}).get("weight_column"),
            "reported_validation_value": baseline.get("reported_value"),
            "resolution_source": source_path,
        }
        if reason:
            spec["warnings"] = [reason]
        return spec

    judged_total = "main_metric" if "main_metric" in getattr(df, "columns", []) else None
    if method == "accuracy" or not columns:
        # Контракт не даёт судимых критериев напрямую (accuracy требует
        # prediction, которого нет в телеметрии). Судим сразу итоговую
        # оценку плана, чтобы не блокировать ассессор ниже по графу.
        if judged_total is not None:
            return _spec(
                judged_total, [], "monitoring_metric_judged_total",
                reason=(
                    f"контракт status={status!r} method={method!r}: критерии "
                    "недоступны напрямую, судится итоговая колонка main_metric"
                ),
            )
        return _spec(
            columns[0] if columns else "target", columns[1:],
            "monitoring_metric_unresolved",
            reason=str(metric.get("reason") or "контракт не computed и в корзине нет main_metric"),
        )
    main = next((name for name in columns if roles.get(name) == "final_score"), columns[0])
    return _spec(main, columns, "monitoring_metric")


def _validate_monitoring_contract_identity(metric: dict, run_context) -> None:
    """Не допустить контракт соседнего агента в текущий запуск."""
    expected = _expected_run_identity(run_context)
    expected_agent = expected.get("agent_ci")
    actual_agent = str(metric.get("basket_id") or "").strip().upper()
    if expected_agent and actual_agent and actual_agent != expected_agent:
        raise ValueError(
            "monitoring_metric принадлежит другому агенту: "
            f"basket_id={actual_agent!r}, run_context.agent_ci={expected_agent!r}"
        )


def _main(
    df: pd.DataFrame,
    assessor_instruction: Path | None = None,
    doc_browser_result: Path | None = None,
    monitoring_metric: dict | None = None,
    model_id: str = "minimax-m2.5",
    llm_model: str = "GigaChat-3-Ultra",
    main_metric: str = "",
    run_context: dict | None = None,
    development_report_artifact=None,
    validation_report_artifact=None,
    model_version_id=0,
):
    """Выбор критериев разметки: ключевая метрика и прочие критерии корзины.

    Сопоставляет бизнес-метрику из инструкции разметки и краткого отчёта о
    разработке с фактическими колонками эталонной корзины. Ответ модели —
    гипотеза: итоговое решение принимает _resolve_main_metric в порядке
    «настройка узла → ответ модели, совпавший с колонкой → единственный
    кандидат → машинный not_computable с причиной». Неоднозначность не
    превращается в случайный выбор и не прерывает невозобновляемый workflow.

    Parameters
    ----------
    df : pandas.DataFrame
        Эталонная корзина (после коннектора эталона).
    assessor_instruction : путь | bytes | dict | None
        Инструкция по разметке.
    doc_browser_result : dict | None
        Карточка отчёта о разработке (ключ bp_card); опциональна.
    model_id : str
        Маршрут запроса: "giga" — GigaChat через langchain, иначе — имя
        модели на AI Gateway (chat/completions).
    llm_model : str
        Имя модели GigaChat для маршрута "giga".
    main_metric : str
        Явная колонка ключевой метрики; приоритетнее ответа модели.
    run_context : dict | None
        Selection manifest текущего прогона. Используется только для сверки
        явно заявленных в отчёте agent/version/distributive; не участвует в
        выборе колонки метрики.
    development_report_artifact : dict | bytes | Path | None
        Исходный DataArtifact отчёта (`bin`/`ext`). Из него извлекаются только
        явные ID для сверки принадлежности и SHA-256; бизнес-смысл продолжает
        приходить из структурированного `doc_browser_result`.
    validation_report_artifact : dict | bytes | Path | None
        Полный отчёт о валидации: raw bytes, локальный путь либо parquet-
        контейнер DataArtifact с bytes/path внутри. Из типизированных таблиц
        извлекаются опубликованные значения КМ, а формула majority vote
        материализуется только когда её пересчёт по корзине воспроизводит
        эти значения. Поданный, но непрочитанный транспорт — техническая
        ошибка ноды, а не статус измерения ``not_computable``.
    model_version_id : int | str
        ID версии из Common Settings; добавляется к ожидаемой identity отчёта.
    """
    validated_contract = None
    if monitoring_metric is not None:
        # wiring v3: селектор — единственный источник контракта для потребителей,
        # поэтому прошедший identity-гейт контракт публикуется на любом пути.
        validated_contract = _load_monitoring_contract(monitoring_metric)
        _validate_monitoring_contract_identity(validated_contract, run_context)
    if validated_contract is not None and not main_metric:
        monitoring_contract = validated_contract
        df = _load_df(df)
        passthrough = _selector_from_monitoring_contract(monitoring_contract, df)
        if passthrough is not None:
            logging.info(
                "kriteria-selector: метрика взята из monitoring_metric "
                "(main=%r, source=%s), LLM и артефакты не используются",
                passthrough["main_metric"], passthrough["resolution_source"],
            )
            return {
                "metric_spec": passthrough,
                "metric_dataset": df,
                "validated_monitoring_metric": monitoring_contract,
            }

    df = _load_df(df)
    metric_df = _metric_rows(df)
    if assessor_instruction is not None:
        instruction_text = _load_instruction(assessor_instruction)
        docx_file = instruction_text
        if len(docx_file) > INSTRUCTION_MAX_CHARS:
            logging.warning(
                "Инструкция длиннее %s символов (%s) — обрезана; проверьте, "
                "что в порт подана инструкция разметки, а не полный отчёт",
                INSTRUCTION_MAX_CHARS, len(docx_file),
            )
            docx_file = docx_file[:INSTRUCTION_MAX_CHARS]
    else:
        instruction_text = ""
        docx_file = ""
    doc_browser_text, document_fields = _doc_browser_context(doc_browser_result)
    validation_evidence = _validation_evidence(
        validation_report_artifact, metric_df
    )
    identity_context = dict(run_context) if isinstance(run_context, dict) else {}
    if model_version_id not in (None, "", 0, "0"):
        identity_context["model_version_id"] = str(model_version_id).strip()
    artifact_provenance = _artifact_provenance(
        df,
        instruction_text,
        doc_browser_result,
        run_context=identity_context,
        development_report_artifact=development_report_artifact,
        validation_report_artifact=validation_report_artifact,
    )
    validation_identity = artifact_provenance["validation_report"]
    if (artifact_provenance["identity_status"] == "mismatch"
            or validation_identity["identity_status"] == "mismatch"):
        resolved = _resolve_main_metric(
            {"main_metric": "", "other_metrics": []},
            metric_df,
            explicit_metric=main_metric,
        )
        json_output = _not_computable_metric_contract(
            {"main_metric": "", "other_metrics": []},
            resolved,
            metric_df,
            document_fields=document_fields,
        )
        mismatches = list(artifact_provenance["identity_mismatches"])
        mismatches.extend(validation_identity["identity_mismatches"])
        json_output.update({
            "reason_code": "artifact_identity_mismatch",
            "reason": (
                "входной отчёт не принадлежит текущему прогону: "
                + "; ".join(mismatches)
            ),
            "validation_evidence": {
                key: value for key, value in validation_evidence.items()
                if key != "text_for_prompt"
            },
            "artifact_provenance": artifact_provenance,
            "selector_inference": {
                "llm_called": False,
                "configured_model_id": model_id,
                "selection_path": "artifact_identity_mismatch",
            },
        })
        logging.error("kriteria-selector: %s", json_output["reason"])
        return {"metric_spec": json_output, "metric_dataset": df}

    # Явная настройка имеет контрактный приоритет и уже проверяется против
    # реальной схемы корзины. Внешний LLM в этой ветке не добавляет информации,
    # зато создаёт задержку, расход квоты и новую точку отказа.
    if str(main_metric or "").strip():
        return _finalize_metric_selection(
            {"main_metric": "", "other_metrics": []},
            metric_df,
            explicit_metric=main_metric,
            document_fields=document_fields,
            artifact_provenance=artifact_provenance,
            validation_evidence=validation_evidence,
            metric_dataset=df,
            model_id=model_id,
        )

    # Сильнейший автоматический путь: формула и опубликованное значение из
    # raw validation report воспроизводятся на raw basket. Здесь LLM не нужен
    # как источник истины; он остаётся fallback для непривычной семантики.
    validation_hypothesis = _majority_hypothesis(
        metric_df, validation_evidence
    ) or _direct_validation_hypothesis(metric_df, validation_evidence)
    if validation_hypothesis is not None:
        metric_dataset, metric_df, materialization_error = (
            _materialize_metric_dataset(
                df, metric_df, validation_hypothesis, validation_evidence
            )
        )
        if materialization_error is None:
            return _finalize_metric_selection(
                validation_hypothesis,
                metric_df,
                explicit_metric="",
                document_fields=document_fields,
                artifact_provenance=artifact_provenance,
                validation_evidence=validation_evidence,
                metric_dataset=metric_dataset,
                model_id=model_id,
            )
        validation_evidence["materialization_error"] = materialization_error
    logging.info(f"model_id: {model_id}, llm_model: {llm_model}")
    system_message = """Ты — инженер измерения качества AI-систем.
Сопоставь ключевую бизнес-метрику с ФАКТИЧЕСКИМИ колонками тестовой корзины.
Не придумывай имена колонок и не вычисляй значение метрики. Если есть
построчная итоговая оценка, выбери её, а не один из частных критериев.
Верни только один JSON-объект, без markdown и без описания JSON-схемы."""

    # Явный список кандидатов сужает ответ модели до реальных колонок:
    # семантическое имя метрики («взвешенная accuracy») вместо колонки —
    # самый частый сбой резолва на корзинах со сводными блоками.
    metric_candidates = _metric_columns(metric_df)
    categorical_candidates = _categorical_metric_columns(metric_df)

    # В корзине формы CI09840650 из проблемного прогона присутствовали только
    # trace/QA-поля и не было ни числовой, ни категориальной разметки. Вызов
    # LLM в этой ситуации ничего не может сопоставить с фактической схемой,
    # поэтому сразу отдаём честный нефатальный отказ и не тратим квоту шлюза.
    if not metric_candidates and not categorical_candidates and not main_metric:
        resolved = _resolve_main_metric(
            {"main_metric": "", "other_metrics": []}, metric_df,
            explicit_metric=main_metric,
        )
        json_output = _not_computable_metric_contract(
            {"main_metric": "", "other_metrics": []}, resolved, metric_df,
            document_fields=document_fields,
        )
        json_output["artifact_provenance"] = artifact_provenance
        logging.warning(
            "kriteria-selector: status=not_computable reason=%s",
            json_output["reason"],
        )
        json_output["validation_evidence"] = {
            key: value for key, value in validation_evidence.items()
            if key != "text_for_prompt"
        }
        json_output["selector_inference"] = {
            "llm_called": False,
            "configured_model_id": model_id,
            "selection_path": "no_metric_candidates",
        }
        return {"metric_spec": json_output, "metric_dataset": df}
    validation_prompt = json.dumps({
        "observations": validation_evidence.get("observations", []),
        "majority_vote_cue": validation_evidence.get("majority_vote_cue", False),
        "mentioned_columns": validation_evidence.get("mentioned_columns", []),
        "relevant_text": validation_evidence.get("text_for_prompt", ""),
    }, ensure_ascii=False)
    user_input = f"""<упрощенный отчёт о разработке>: {doc_browser_text}\n\n

    <Проверенные данные полного отчёта о валидации>: {validation_prompt}\n\n

    ---------

    <Пример данных из датасета:

    Data Type: {metric_df.dtypes}

    n.unique: {metric_df.nunique()}

    df head(2): {_truncate_cells(metric_df.head(2))}

    ------

    <Инструкция по разметке, в которой содержатся критерии разметки>

    {docx_file}

    ----

    <КРАЙНЕ ВАЖНО> Верни один JSON-объект следующей формы.
    Все названия колонок перепиши дословно из Data Type.
    Числовые колонки-кандидаты в ключевую метрику: {metric_candidates}.
    Категориальные колонки разметки (при отсутствии числовой шкалы):
    {categorical_candidates}.
    В main_metric подставь ровно одно имя из этих списков — НЕ название
    бизнес-метрики («accuracy», «взвешенная accuracy»), а имя колонки:

    {{
        "main_metric": "имя итоговой колонки разметки",
        "metric_name": "семантическое название бизнес-метрики",
        "other_metrics": ["колонки частных критериев"],
        "source_columns": ["колонки голосов, если итог вычисляется"],
        "row_aggregation": "identity",
        "weight_column": null,
        "prediction_column": null,
        "reference_column": null,
        "strategy": "mean_score",
        "scale_min": null,
        "scale_max": null,
        "higher_is_better": true
    }}

    strategy: mean_score для готовой построчной итоговой оценки;
    weighted_accuracy только если есть отдельная колонка веса/частоты;
    accuracy только для пары prediction/reference без готовой оценки.
    Если отчёт явно задаёт мнение большинства, верни
    row_aggregation=majority_vote и все source_columns голосов; не выбирай
    ошибочную готовую финальную колонку. Во всех остальных случаях используй
    row_aggregation=identity и source_columns=[main_metric].

    """

    from langchain_core.output_parsers import JsonOutputParser
    from langchain_core.output_parsers.string import StrOutputParser

    fails = 0
    last_error = None
    json_output = None

    while json_output is None and fails < LLM_ATTEMPTS:
        try:
            if model_id != "giga":
                response = make_sds_request(
                    model_id=model_id,
                    input_prompt=user_input,
                    system_prompt=system_message,
                )
                json_output = _parse_metrics_json(_chat_content(response))
            else:
                response = make_giga_request(
                    input_prompt=system_message + "\n\n" + user_input,
                    model=llm_model,
                )
                str_output = StrOutputParser().invoke(response)
                try:
                    json_output = JsonOutputParser().invoke(str_output)
                except Exception:
                    json_output = _parse_metrics_json(str_output)  # снять ```json```/текст

        except Exception as e:
            fails += 1
            last_error = e
            logging.exception(
                "Ошибка выбора критериев, попытка %s/%s", fails, LLM_ATTEMPTS
            )
            if fails < LLM_ATTEMPTS:
                # Экспоненциальная пауза: немедленный повтор попадает в тот же
                # всплеск нагрузки шлюза и сгорает впустую.
                time.sleep(RETRY_BASE_DELAY_SECONDS * 2 ** (fails - 1))

    if json_output is None:
        # LLM не ответил валидным JSON. Не выдумываем метрики за него:
        # решение примет _resolve_main_metric (явная настройка / единственный
        # кандидат / понятная остановка).
        logging.warning(
            "kriteria-selector: модель не вернула валидный JSON (%s)", last_error
        )
        json_output = {"main_metric": "", "other_metrics": []}

    if not isinstance(json_output, dict):
        logging.warning(
            "kriteria-selector: ответ модели должен быть dict, получено %s; "
            "используется безопасный резолв по схеме корзины",
            type(json_output),
        )
        json_output = {"main_metric": "", "other_metrics": []}
    if "other_metrics" not in json_output or json_output["other_metrics"] is None:
        json_output["other_metrics"] = []
    if not isinstance(json_output["other_metrics"], list):
        logging.warning(
            "kriteria-selector: other_metrics должен быть list; значение %r "
            "отброшено", json_output["other_metrics"],
        )
        json_output["other_metrics"] = []

    metric_dataset, metric_df, materialization_error = _materialize_metric_dataset(
        df, metric_df, json_output, validation_evidence
    )
    if materialization_error:
        # Неисполнимая формула агрегации не отменяет сам выбор критериев:
        # прежний not_computable c main_metric='target' блокировал ассессор
        # ниже по графу. Судья размечает исходные колонки (identity), причина
        # сохраняется предупреждением в evidence.
        logging.warning(
            "kriteria-selector: формула не материализована (%s); агрегация "
            "понижена до identity, критерии сохранены", materialization_error,
        )
        validation_evidence["materialization_error"] = materialization_error
        json_output = dict(json_output)
        json_output["row_aggregation"] = "identity"

    # Итоговое решение по метрике: ответ LLM → единственный кандидат.
    # При неоднозначности — not_computable с перечнем колонок, без угадывания.
    return _finalize_metric_selection(
        json_output,
        metric_df,
        explicit_metric=main_metric,
        document_fields=document_fields,
        artifact_provenance=artifact_provenance,
        validation_evidence=validation_evidence,
        metric_dataset=metric_dataset,
        model_id=model_id,
        llm_called=True,
    )


def main(*args, **kwargs):
    """Обёртка контура: доложить validated_monitoring_metric на любом пути.

    Контракт публикуется, только если он был подан и прошёл identity-гейт
    (иначе _main поднял ValueError до этой точки).
    """
    monitoring_metric = kwargs.get("monitoring_metric")
    run_context = kwargs.get("run_context")
    result = _main(*args, **kwargs)
    if monitoring_metric is not None and "validated_monitoring_metric" not in result:
        contract = _load_monitoring_contract(monitoring_metric)
        _validate_monitoring_contract_identity(contract, run_context)
        result["validated_monitoring_metric"] = contract
    return result
