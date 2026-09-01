"""MetricSpec обязан быть прослеживаем до входных артефактов."""

import importlib.util
import io
import json
import tempfile
from pathlib import Path

import pandas as pd
from docx import Document


MODULE_DIR = Path(__file__).resolve().parents[1]
SPEC = importlib.util.spec_from_file_location(
    "selector_artifact_provenance", MODULE_DIR / "main.py"
)
selector = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(selector)


def test_validation_report_port_is_mounted_as_local_path():
    descriptor = json.loads((MODULE_DIR / "descriptor.json").read_text())
    port = next(
        item for item in descriptor["ports"]
        if item["name"] == "validation_report_artifact"
    )
    assert port["getPortAsLocalPath"] is True


def test_monitoring_metric_path_bypasses_legacy_and_keeps_accuracy_total(tmp_path):
    frame = pd.DataFrame({
        "input_query": ["q1", "q2"],
        "output_answer": ["a1", "a2"],
        "class": ["x", "y"],
        "GT": ["x", "x"],
        "main_metric": [1.0, 0.0],
    })
    contract = {
        "contract_version": "laim-monitoring-metric.v2",
        "status": "computed",
        "name": "Accuracy",
        "score_column": "main_metric",
        "assessment_mode": "qa",
        "scoring": {
            "method": "accuracy",
            "sources": [
                {"column_name": "class", "role": "prediction"},
                {"column_name": "GT", "role": "target"},
            ],
        },
        "aggregation": {"method": "mean"},
        "baseline": {"reported_value": 0.97},
    }
    path = tmp_path / "tmp.monitoring_metric"
    path.write_text(json.dumps(contract), encoding="utf-8")

    output = selector.main(
        df=frame,
        monitoring_metric=str(path),
        validation_report_artifact={"bin": b"not a docx"},
    )

    spec = output["metric_spec"]
    assert spec["main_metric"] == "main_metric"
    assert spec["resolution_source"] == "monitoring_metric_judged_total"
    assert output["metric_dataset"].equals(frame)


def test_monitoring_majority_exposes_every_vote_to_assessor():
    frame = pd.DataFrame({
        "input_query": ["q1", "q2"],
        "output_answer": ["a1", "a2"],
        "mark1": [1, 0],
        "mark2": [1, 0],
        "mark3": [0, 1],
        "main_metric": [1.0, 0.0],
    })
    contract = {
        "contract_version": "laim-monitoring-metric.v2",
        "status": "computed",
        "name": "Accuracy",
        "score_column": "main_metric",
        "assessment_mode": "dialogue",
        "scoring": {
            "method": "majority",
            "sources": [
                {"column_name": "mark1", "role": "assessor_vote"},
                {"column_name": "mark2", "role": "assessor_vote"},
                {"column_name": "mark3", "role": "assessor_vote"},
            ],
        },
        "aggregation": {"method": "mean"},
    }

    spec = selector.main(
        df=frame,
        monitoring_metric={"monitoring_metric": contract},
    )["metric_spec"]

    assert spec["main_metric"] == "mark1"
    assert spec["other_metrics"] == ["mark2", "mark3"]
    assert spec["source_columns"] == ["mark1", "mark2", "mark3"]
    assert spec["row_aggregation"] == "majority_vote"
    assert spec["assessment_mode"] == "dialogue"


def test_monitoring_metric_of_another_agent_fails_before_selection():
    frame = pd.DataFrame({
        "input_query": ["q"], "output_answer": ["a"], "main_metric": [1]
    })
    contract = {
        "basket_id": "CI09997554",
        "status": "computed",
        "scoring": {"method": "identity", "sources": [
            {"column_name": "main_metric", "role": "final_score"}
        ]},
    }
    try:
        selector.main(
            df=frame,
            monitoring_metric=contract,
            run_context={"agent_ci": "CI09997438"},
        )
    except ValueError as error:
        assert "другому агенту" in str(error)
        assert "CI09997554" in str(error)
        assert "CI09997438" in str(error)
    else:
        raise AssertionError("selector принял monitoring_metric соседнего агента")


def test_provenance_is_stable_and_does_not_expose_artifact_text():
    basket = pd.DataFrame(
        {
            "input_query": ["секретный вопрос", "другой вопрос"],
            "output_answer": ["секретный ответ", "другой ответ"],
            "Итог": [1, 0],
        }
    )
    report = {
        "extracted_fields": {
            "bp_card_fields": {
                "evaluation_metric": "Accuracy",
                "threshold": "80%",
            }
        }
    }

    first = selector._artifact_provenance(
        basket, "секретная инструкция", report
    )
    second = selector._artifact_provenance(
        basket.copy(), "секретная инструкция", report
    )

    assert first == second
    assert first["basket"]["sha256"]
    assert first["basket"]["n_rows"] == 2
    assert first["instruction"]["sha256"]
    assert first["development_report"]["sha256"]
    assert first["identity_status"] == "not_provided"
    serialized = str(first)
    assert "секретный вопрос" not in serialized
    assert "секретная инструкция" not in serialized


def test_artifact_identity_is_verified_against_run_context():
    frame = pd.DataFrame({"input_query": ["q"], "output_answer": ["a"], "target": [1]})
    report = {
        "summary": (
            "Дистрибутив CI09840670-D-01.004.02-604-distrib.zip; "
            "ID версии модели в БМ | 342293"
        )
    }
    provenance = selector._artifact_provenance(
        frame,
        "инструкция",
        report,
        run_context={"agent_ci": "CI09840670", "distributive": "D-01.004.02-604"},
    )
    assert provenance["identity_status"] == "verified"
    assert provenance["expected_identity"]["agent_ci"] == "CI09840670"
    assert "CI09840670" in provenance["declared_identity"]["solution_ids"]
    assert "342293" in provenance["declared_identity"]["model_version_ids"]


def test_artifact_identity_mismatch_is_explicit():
    frame = pd.DataFrame({"input_query": ["q"], "output_answer": ["a"], "target": [1]})
    provenance = selector._artifact_provenance(
        frame,
        "инструкция",
        {"summary": "Дистрибутив CI09877398-D-01.006.00-515-distrib.zip"},
        run_context={"agent_ci": "CI09840670", "distributive": "D-01.004.02-604"},
    )
    assert provenance["identity_status"] == "mismatch"
    assert provenance["identity_mismatches"]


def test_identity_mismatch_refuses_metric_before_llm_call():
    frame = pd.DataFrame(
        {"input_query": ["q"], "output_answer": ["a"], "target": [1]}
    )
    result = selector.main(
        df=frame,
        assessor_instruction={"text": "оцените target"},
        doc_browser_result={"summary": "Решение CI09877398"},
        run_context={"agent_ci": "CI09840670"},
    )["metric_spec"]
    assert result["status"] == "not_computable"
    assert result["reason_code"] == "artifact_identity_mismatch"
    assert result["artifact_provenance"]["identity_status"] == "mismatch"


def test_raw_development_report_bytes_supply_verifiable_identity():
    document = Document()
    document.add_paragraph("ID версии модели в БМ | 342293")
    document.add_paragraph(
        "https://nexus/CI07909419/CI09840670/D-01.004.02-604/"
        "CI09840670-D-01.004.02-604-distrib.zip"
    )
    buffer = io.BytesIO()
    document.save(buffer)
    frame = pd.DataFrame({"input_query": ["q"], "output_answer": ["a"], "target": [1]})

    provenance = selector._artifact_provenance(
        frame,
        "инструкция",
        {"bp_card_fields": {"evaluation_metric": "Accuracy"}},
        run_context={"agent_ci": "CI09840670", "distributive": "D-01.004.02-604"},
        development_report_artifact={"bin": buffer.getvalue(), "ext": ".docx"},
    )
    assert provenance["identity_status"] == "verified"
    assert provenance["development_report"]["artifact_sha256"]
    assert provenance["development_report"]["artifact_n_bytes"] > 0


def test_validation_header_uses_second_number_as_model_version():
    document = Document()
    document.add_paragraph(
        "ID GenAI-решения / ID версии GenAI-решения: 300532 / 351335"
    )
    buffer = io.BytesIO()
    document.save(buffer)
    frame = pd.DataFrame({"input_query": ["q"], "output_answer": ["a"], "target": [1]})

    provenance = selector._artifact_provenance(
        frame,
        "инструкция",
        None,
        run_context={"model_version_id": "351335"},
        validation_report_artifact={"bin": buffer.getvalue(), "ext": ".docx"},
    )

    validation = provenance["validation_report"]
    assert validation["identity_status"] == "verified"
    assert validation["identity_verified_fields"] == ["model_version_id"]


def test_long_structured_report_text_is_not_treated_as_filesystem_path():
    raw, extension = selector._development_report_bytes("x" * 10000)

    assert raw is None
    assert extension == ""


def test_explicit_metric_does_not_call_external_llm():
    frame = pd.DataFrame(
        {"input_query": ["q"], "output_answer": ["a"], "target": [1]}
    )
    original = selector.make_sds_request
    selector.make_sds_request = lambda *args, **kwargs: (_ for _ in ()).throw(
        AssertionError("LLM не должен вызываться при явной метрике")
    )
    try:
        result = selector.main(
            df=frame,
            assessor_instruction={"text": "реальная инструкция"},
            doc_browser_result=None,
            main_metric="target",
        )["metric_spec"]
    finally:
        selector.make_sds_request = original

    assert result["status"] == "resolved"
    assert result["main_metric"] == "target"
    assert result["resolution_source"] == "setting"


def _validation_docx(rows, paragraph=""):
    document = Document()
    if paragraph:
        document.add_paragraph(paragraph)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Наименование критерия"
    table.rows[0].cells[1].text = "Значение"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return {"bin": buffer.getvalue(), "ext": ".docx"}


def _parquet_transport(artifact):
    """Транспорт dataframe-порта SberDS: один DataArtifact внутри parquet."""
    buffer = io.BytesIO()
    pd.DataFrame({
        "bin": [artifact["bin"]],
        "ext": [artifact["ext"]],
    }).to_parquet(buffer, index=False)
    return buffer.getvalue()


def test_validation_report_is_unwrapped_from_parquet_bytes_and_local_path():
    artifact = _validation_docx([("Accuracy", "0.90")])
    parquet_bytes = _parquet_transport(artifact)

    text, tables, error = selector._validation_report_content(parquet_bytes)
    assert error is None
    assert "Accuracy" in text
    assert len(tables) == 1

    with tempfile.TemporaryDirectory() as directory:
        # getPortAsLocalPath в контуре часто возвращает имя без расширения.
        local_path = Path(directory) / "validation_report_artifact"
        local_path.write_bytes(parquet_bytes)
        path_text, path_tables, path_error = selector._validation_report_content(
            {"path": str(local_path)}
        )
    assert path_error is None
    assert path_text == text
    assert path_tables == tables


def test_unstaged_hdfs_uri_is_a_technical_error_not_missing_evidence():
    try:
        selector._validation_evidence(
            "hdfs:///user/example/validation_report.docx",
            pd.DataFrame({"target": [1]}),
        )
    except selector.ArtifactTransportError as error:
        assert "получен только HDFS URI" in str(error)
        assert "локально смонтированный" in str(error)
    else:
        raise AssertionError("HDFS URI без смонтированного файла принят как отчёт")


def test_validation_report_materializes_generic_majority_and_weight():
    majority = [1] * 15 + [0] * 5
    frame = pd.DataFrame({
        "input_query": [f"q{i}" for i in range(20)],
        "output_answer": [f"a{i}" for i in range(20)],
        "reviewer_a": majority,
        "reviewer_b": majority,
        "reviewer_c": [1 - value for value in majority],
        "sample_weight": [1] * 15 + [2] * 5,
        "final_score": [0] * 20,
    })
    validation = _validation_docx(
        [("Accuracy (macro)", "0.75"), ("Accuracy (weighted)", "0.60")],
        "Итоговый результат определялся большинством голосов "
        "reviewer_a, reviewer_b и reviewer_c.",
    )
    original = selector.make_sds_request
    selector.make_sds_request = lambda *args, **kwargs: (_ for _ in ()).throw(
        AssertionError("доказанная формула не должна зависеть от LLM")
    )
    try:
        output = selector.main(
            df=frame,
            assessor_instruction={"text": "бинарная оценка"},
            doc_browser_result=None,
            validation_report_artifact=validation,
        )
    finally:
        selector.make_sds_request = original

    spec = output["metric_spec"]
    assert spec["status"] == "resolved"
    assert spec["row_aggregation"] == "majority_vote"
    assert spec["source_columns"] == ["reviewer_a", "reviewer_b", "reviewer_c"]
    assert spec["weight_column"] == "sample_weight"
    assert spec["validation_evidence"]["status"] == "confirmed"
    assert output["metric_dataset"]["laim_key_metric"].tolist() == majority


def test_validation_contradiction_is_nonfatal_not_computable():
    frame = pd.DataFrame({
        "input_query": ["q1", "q2", "q3", "q4"],
        "output_answer": ["a1", "a2", "a3", "a4"],
        "target": [1, 1, 1, 0],
    })
    output = selector.main(
        df=frame,
        assessor_instruction={"text": "бинарная оценка"},
        doc_browser_result=None,
        validation_report_artifact=_validation_docx(
            [("Accuracy", "0.10")]
        ),
        main_metric="target",
    )
    spec = output["metric_spec"]
    assert spec["status"] == "not_computable"
    assert spec["reason_code"] == "validation_metric_not_reproduced"
    assert spec["validation_evidence"]["status"] == "contradiction"
    assert output["metric_dataset"].equals(frame)


def test_unreadable_validation_report_fails_node_with_transport_reason():
    frame = pd.DataFrame({
        "input_query": ["q"], "output_answer": ["a"], "target": [1]
    })
    try:
        selector.main(
            df=frame,
            assessor_instruction={"text": "бинарная оценка"},
            doc_browser_result=None,
            validation_report_artifact={"bin": b"not a docx", "ext": ".docx"},
            main_metric="target",
        )
    except selector.ArtifactTransportError as error:
        assert "validation_report_artifact не прочитан" in str(error)
        assert "DOCX не прочитан" in str(error)
    else:
        raise AssertionError("непрочитанный validation report не остановил ноду")


def test_monitoring_metric_passthrough_publishes_validated_contract():
    """Контракт, прошедший identity-гейт, уходит потребителям без изменений."""
    frame = pd.DataFrame({
        "input_query": ["q1"], "output_answer": ["a1"], "quality_metric": [1.0],
    })
    contract = {
        "contract_version": "laim-monitoring-metric.v2",
        "status": "computed",
        "basket_id": "CI09999999",
        "name": "quality",
        "assessment_mode": "qa",
        "scoring": {
            "method": "identity",
            "sources": [{
                "source_id": "source_1", "column_name": "quality_metric",
                "role": "final_score", "normalization": "numeric", "polarity": "direct",
            }],
        },
        "baseline": {"value": 0.9},
    }

    result = selector.main(
        df=frame,
        monitoring_metric=contract,
        run_context={"selection": {"agent_ci": "CI09999999"}},
    )

    assert result["validated_monitoring_metric"] == contract
    assert result["metric_spec"]["status"] == "resolved"


def test_packed_dialogue_metadata_is_not_a_metric_candidate():
    frame = pd.DataFrame({
        "session_id": ["s1", "s2"],
        "dialogue": ["[(1, 'q', 'a')]", "[(2, 'q2', 'a2')]"],
        "input_query_count": [1, 2],
        "turn_index": [1, 1],
        "assessment_unit_id": ["s1", "s2"],
        "solution_version": ["v1", "v1"],
        "quality_metric": [1, 0],
    })

    assert selector._metric_columns(frame) == ["quality_metric"]


def test_validated_contract_is_published_even_on_manual_override_path():
    """wiring v3: селектор — единственный источник контракта для потребителей,
    поэтому прошедший identity-гейт контракт публикуется и на legacy-пути."""
    frame = pd.DataFrame({
        "input_query": ["q1"], "output_answer": ["a1"], "quality_metric": [1.0],
    })
    contract = {
        "contract_version": "laim-monitoring-metric.v2",
        "status": "computed",
        "basket_id": "CI09999999",
        "assessment_mode": "qa",
        "scoring": {
            "method": "identity",
            "sources": [{
                "source_id": "source_1", "column_name": "quality_metric",
                "role": "final_score", "normalization": "numeric", "polarity": "direct",
            }],
        },
    }

    result = selector.main(
        df=frame,
        monitoring_metric=contract,
        main_metric="quality_metric",  # ручное переопределение уводит с passthrough
        run_context={"selection": {"agent_ci": "CI09999999"}},
    )

    assert result["validated_monitoring_metric"] == contract
    assert result["metric_spec"]["main_metric"] == "quality_metric"
